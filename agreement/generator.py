"""
Codesino Agreement Generator — Python / python-docx implementation
===================================================================
File: agreement/generator.py

Four generator functions, one per document type:
  • generate_agreement_buffer(data)            → bytes  [Client Project Agreement]
  • generate_executive_agreement_buffer(data)  → bytes  [Executive / Senior Roles]
  • generate_support_agreement_buffer(data)    → bytes  [Support / Operations Roles]
  • generate_technical_agreement_buffer(data)  → bytes  [Technical / Dev Team Roles]

All return raw .docx bytes ready to stream as an HTTP response.
"""

import os
import random
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.style import WD_STYLE_TYPE

# ─────────────────────────────────────────────────────────────────────────────
# BRAND COLOURS  —  Professional Blue Palette
# ─────────────────────────────────────────────────────────────────────────────
BRAND_BLUE   = RGBColor(0x1A, 0x4A, 0x8A)
MID_BLUE     = RGBColor(0x25, 0x6B, 0xC4)
LIGHT_BLUE   = RGBColor(0xE8, 0xF1, 0xFB)
PALE_BLUE    = RGBColor(0xD0, 0xE4, 0xF7)
DARK_GRAY    = RGBColor(0x22, 0x22, 0x22)
MID_GRAY     = RGBColor(0x55, 0x55, 0x55)
WHITE_COLOR  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_ACCENT = RGBColor(0x19, 0x7A, 0x4A)
PLACEHOLDER  = RGBColor(0xAA, 0xAA, 0xAA)
MUTED_BLUE   = RGBColor(0x5A, 0x7E, 0xB5)

HEX_BRAND_BLUE = "1A4A8A"
HEX_MID_BLUE   = "256BC4"
HEX_LIGHT_BLUE = "E8F1FB"
HEX_PALE_BLUE  = "D0E4F7"
HEX_DARK_GRAY  = "222222"
HEX_WHITE      = "FFFFFF"
HEX_RULE       = "B8CDE8"
HEX_GREEN      = "197A4A"
HEX_FAFAFA     = "FAFAFA"
HEX_F0F8FF     = "F0F8FF"   # faint blue for selected option box (client agreement)

# ─────────────────────────────────────────────────────────────────────────────
# IMAGE PATHS
# ─────────────────────────────────────────────────────────────────────────────
_HERE      = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH  = os.path.join(_HERE, "images", "codesino_logo.png")
STAMP_PATH = os.path.join(_HERE, "images", "codesino_stamp.png")

# ─────────────────────────────────────────────────────────────────────────────
# POSITION CATALOGUES  (employment agreements)
# ─────────────────────────────────────────────────────────────────────────────
EXECUTIVE_POSITIONS = [
    "Chief Executive Officer (CEO)",
    "Chief Technology Officer (CTO)",
    "Chief Operating Officer (COO)",
    "Chief Financial Officer (CFO)",
    "Chief Marketing Officer (CMO)",
    "Co-Founder",
    "Managing Director",
    "General Manager",
    "Head of Engineering",
    "Head of Product",
    "Head of Design",
    "Head of Business Development",
    "Director of Operations",
    "Director of Finance",
    "Director of Marketing",
    "Senior Project Manager",
    "Team Lead",
    "Other (Specified)",
]

SUPPORT_POSITIONS = [
    "Customer Support Representative",
    "Customer Success Manager",
    "Social Media Manager",
    "Social Media Content Creator",
    "Community Manager",
    "Marketing Associate",
    "Sales Representative",
    "Administrative Assistant",
    "Accounts & Billing Officer",
    "Quality Assurance (QA) Analyst",
    "Data Entry Specialist",
    "HR / Recruitment Officer",
    "Brand Ambassador",
    "Virtual Assistant",
    "Other (Specified)",
]

TECHNICAL_POSITIONS = [
    "Frontend Developer",
    "Backend Developer",
    "Full Stack Developer",
    "Mobile Developer (React Native / Flutter)",
    "UI/UX Designer",
    "DevOps Engineer",
    "Cloud Infrastructure Engineer",
    "Database Administrator",
    "Cybersecurity Engineer",
    "Machine Learning Engineer",
    "API Integration Specialist",
    "Software Architect",
    "Technical Project Manager",
    "QA / Test Automation Engineer",
    "Other (Specified)",
]

EMPLOYMENT_TYPES = [
    "Full-Time",
    "Part-Time",
    "Contract",
    "Internship",
    "Freelance / Retainer",
    "Probationary",
]

SALARY_TYPES = [
    "Monthly Fixed Salary",
    "Percentage Per Contract",
    "Hourly Rate",
    "Commission-Based",
    "Retainer Fee",
]

# ─────────────────────────────────────────────────────────────────────────────
# LOW-LEVEL XML HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    for ex in tcPr.findall(qn("w:shd")):
        tcPr.remove(ex)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tbl_borders = OxmlElement("w:tcBorders")
    for side, cfg in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if cfg is None:
            el.set(qn("w:val"), "none")
        else:
            el.set(qn("w:val"),   cfg.get("val",   "single"))
            el.set(qn("w:sz"),    str(cfg.get("sz",  4)))
            el.set(qn("w:color"), cfg.get("color",  "000000"))
        tbl_borders.append(el)
    for ex in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(ex)
    tcPr.append(tbl_borders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    for ex in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(ex)
    tcPr.append(mar)


def set_cell_vertical_align(cell, align="center"):
    tc     = cell._tc
    tcPr   = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    for ex in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(ex)
    tcPr.append(vAlign)


def set_cell_width(cell, width_dxa):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    for ex in tcPr.findall(qn("w:tcW")):
        tcPr.remove(ex)
    tcPr.append(tcW)


def add_paragraph_border_bottom(para, color=HEX_RULE, sz=6, space=4):
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    for ex in pPr.findall(qn("w:pBdr")):
        pPr.remove(ex)
    pPr.append(pBdr)


def add_paragraph_spacing(para, before=0, after=0, line=None):
    pPr     = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    if before is not None:
        spacing.set(qn("w:before"), str(before))
    if after is not None:
        spacing.set(qn("w:after"),  str(after))
    if line is not None:
        spacing.set(qn("w:line"),     str(line))
        spacing.set(qn("w:lineRule"), "auto")
    for ex in pPr.findall(qn("w:spacing")):
        pPr.remove(ex)
    pPr.append(spacing)


def set_run_font(run, font_name="Calibri"):
    rPr    = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    for ex in rPr.findall(qn("w:rFonts")):
        rPr.remove(ex)
    rPr.insert(0, rFonts)


def set_table_width(table, width_dxa):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    for ex in tblPr.findall(qn("w:tblW")):
        tblPr.remove(ex)
    tblPr.append(tblW)


def set_column_widths(table, widths):
    tbl    = table._tbl
    for ex in tbl.findall(qn("w:tblGrid")):
        tbl.remove(ex)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def set_table_no_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    for ex in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(ex)
    tblPr.append(tblBorders)


def no_space(para):
    add_paragraph_spacing(para, before=0, after=0)


def clear_cell_paras(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


# ─────────────────────────────────────────────────────────────────────────────
# MISC HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def generate_ref(prefix="EMP"):
    now  = datetime.now()
    y    = str(now.year)[-2:]
    m    = str(now.month).zfill(2)
    d    = str(now.day).zfill(2)
    rand = random.randint(1000, 9999)
    return f"CD-{prefix}-{y}{m}{d}-{rand}"


def format_date(date_str=None):
    if not date_str:
        return datetime.now().strftime("%d %B %Y")
    try:
        dt = datetime.strptime(str(date_str)[:10], "%Y-%m-%d")
        return dt.strftime("%d %B %Y")
    except Exception:
        return str(date_str)


def format_datetime():
    return datetime.now().strftime("%d %B %Y, %H:%M")


def cb(checked):
    return "☑" if checked else "☐"


def v(val, fallback="Not Provided"):
    return val if val else fallback


# Backwards-compat alias (old client-agreement generator used val_or)
def val_or(val, fallback="Not Provided"):
    return val if val else fallback


BODY_FONT = "Calibri"

# ─────────────────────────────────────────────────────────────────────────────
# TYPOGRAPHY PRIMITIVES
# ─────────────────────────────────────────────────────────────────────────────

def add_run(para, text, bold=False, italic=False, size_pt=11,
            color=None, font=BODY_FONT):
    run = para.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size_pt)
    run.font.name   = font
    if color:
        run.font.color.rgb = color if isinstance(color, RGBColor) else hex_to_rgb(color)
    set_run_font(run, font)
    return run


def heading1(doc, text):
    para = doc.add_paragraph()
    add_paragraph_spacing(para, before=300, after=80)
    add_paragraph_border_bottom(para, color=HEX_BRAND_BLUE, sz=10, space=4)
    add_run(para, text, bold=True, size_pt=13, color=BRAND_BLUE)
    return para


def heading2(doc, text):
    para = doc.add_paragraph()
    add_paragraph_spacing(para, before=160, after=60)
    add_run(para, text, bold=True, size_pt=11, color=MID_BLUE)
    return para


def body_text(doc, text, italic=False, color=None):
    para = doc.add_paragraph()
    add_paragraph_spacing(para, before=40, after=60, line=276)
    add_run(para, text, italic=italic, size_pt=11,
            color=color if color else DARK_GRAY)
    return para


def italic_note(doc, text):
    para = doc.add_paragraph()
    add_paragraph_spacing(para, before=30, after=60)
    add_run(para, text, italic=True, size_pt=10, color=MUTED_BLUE)
    return para


def spacer(doc, before=80, after=80):
    para = doc.add_paragraph()
    add_paragraph_spacing(para, before=before, after=after)
    return para


def divider(doc):
    para = doc.add_paragraph()
    add_paragraph_spacing(para, before=180, after=180)
    add_paragraph_border_bottom(para, color=HEX_RULE, sz=2, space=1)
    return para


def lv_para(container, label, value, label_color=None):
    if label_color is None:
        label_color = BRAND_BLUE
    para = container.add_paragraph()
    add_paragraph_spacing(para, before=0, after=56)
    r1 = para.add_run(f"{label}:  ")
    r1.bold             = True
    r1.font.size        = Pt(11)
    r1.font.color.rgb   = label_color
    r1.font.name        = BODY_FONT
    set_run_font(r1)
    is_empty = not value or str(value).strip() == ""
    r2 = para.add_run("Not Provided" if is_empty else str(value))
    r2.font.size      = Pt(11)
    r2.italic         = is_empty
    r2.font.color.rgb = PLACEHOLDER if is_empty else DARK_GRAY
    r2.font.name      = BODY_FONT
    set_run_font(r2)
    return para


def section_box(doc, fill_fn, fill_hex=HEX_WHITE, accent_color=HEX_BRAND_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table,    9360)
    set_column_widths(table, [9360])
    set_table_no_borders(table)
    cell = table.cell(0, 0)
    set_cell_width(cell,    9360)
    set_cell_shading(cell,  fill_hex)
    set_cell_margins(cell,  top=140, bottom=140, left=220, right=160)
    set_cell_borders(
        cell,
        top    = {"val": "single", "sz": 2,  "color": HEX_RULE},
        bottom = {"val": "single", "sz": 2,  "color": HEX_RULE},
        left   = {"val": "single", "sz": 18, "color": accent_color},
        right  = None,
    )
    clear_cell_paras(cell)
    fill_fn(cell)
    return table


def data_table(doc, headers, rows, col_widths):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    set_table_width(table,    sum(col_widths))
    set_column_widths(table,  col_widths)
    set_table_no_borders(table)
    thin = {"val": "single", "sz": 2, "color": HEX_RULE}
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_width(cell,   col_widths[i])
        set_cell_shading(cell, HEX_BRAND_BLUE)
        set_cell_margins(cell, top=100, bottom=100, left=140, right=120)
        set_cell_borders(cell, top=thin, bottom=thin, left=thin, right=thin)
        clear_cell_paras(cell)
        p = cell.add_paragraph()
        no_space(p)
        add_run(p, h, bold=True, size_pt=10, color=WHITE_COLOR)
    for ri, row_data in enumerate(rows):
        row  = table.rows[ri + 1]
        fill = HEX_WHITE if ri % 2 == 0 else HEX_LIGHT_BLUE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_width(cell,   col_widths[ci])
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=90, bottom=90, left=140, right=120)
            set_cell_borders(cell, top=thin, bottom=thin, left=thin, right=thin)
            clear_cell_paras(cell)
            p = cell.add_paragraph()
            no_space(p)
            add_run(p, str(val) if val is not None else "", size_pt=10, color=DARK_GRAY)
    return table


def add_bullet(container, text):
    try:
        para = container.add_paragraph(style="List Bullet")
    except Exception:
        para = container.add_paragraph()
    add_paragraph_spacing(para, before=30, after=30)
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    add_run(para, text, size_pt=11, color=DARK_GRAY)
    return para


def add_numbered(container, text):
    try:
        para = container.add_paragraph(style="List Number")
    except Exception:
        para = container.add_paragraph()
    add_paragraph_spacing(para, before=30, after=30)
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    add_run(para, text, size_pt=11, color=DARK_GRAY)
    return para


def sig_line(cell, label="Signature:"):
    p1 = cell.add_paragraph()
    add_paragraph_spacing(p1, before=0, after=20)
    add_run(p1, label, size_pt=10, color=MID_GRAY)
    for _ in range(2):
        pb = cell.add_paragraph()
        no_space(pb)
    pu = cell.add_paragraph()
    add_paragraph_spacing(pu, before=0, after=60)
    add_paragraph_border_bottom(pu, color=HEX_BRAND_BLUE, sz=4, space=1)


def add_image_para(container, path, width_inches, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add an image paragraph; silently skips if file is missing."""
    if not os.path.exists(path):
        return None
    para = container.add_paragraph()
    no_space(para)
    para.alignment = align
    run  = para.add_run()
    run.add_picture(path, width=Inches(width_inches))
    return para


# ─────────────────────────────────────────────────────────────────────────────
# SHARED: DOCUMENT SETUP
# ─────────────────────────────────────────────────────────────────────────────

def _setup_document():
    doc = Document()
    for section in doc.sections:
        section.page_width    = Twips(12240)
        section.page_height   = Twips(15840)
        section.left_margin   = Twips(1440)
        section.right_margin  = Twips(1440)
        section.top_margin    = Twips(1080)
        section.bottom_margin = Twips(1080)
    for style_name, style_type in [
        ("List Bullet", WD_STYLE_TYPE.PARAGRAPH),
        ("List Number", WD_STYLE_TYPE.PARAGRAPH),
    ]:
        try:
            doc.styles[style_name]
        except KeyError:
            doc.styles.add_style(style_name, style_type)
    return doc


# ─────────────────────────────────────────────────────────────────────────────
# SHARED: COVER HEADER  (logo inside the blue title band) — employment agreements
# ─────────────────────────────────────────────────────────────────────────────

def _build_header(doc, subtitle, agreement_ref, agreement_date):
    """
    Two-column header:
      Left  → deep-blue band with logo image centred inside it
      Right → deeper-blue band with document title + subtitle
    Then a light-blue ref/date bar beneath.
    """
    hdr_table = doc.add_table(rows=1, cols=2)
    set_table_width(hdr_table,    9360)
    set_column_widths(hdr_table, [2200, 7160])
    set_table_no_borders(hdr_table)

    logo_cell  = hdr_table.cell(0, 0)
    title_cell = hdr_table.cell(0, 1)

    # ── Logo cell: blue background, logo centred inside ──────────────────────
    set_cell_width(logo_cell,    2200)
    set_cell_shading(logo_cell,  HEX_BRAND_BLUE)
    set_cell_margins(logo_cell,  top=140, bottom=140, left=100, right=80)
    set_cell_borders(logo_cell,  top=None, bottom=None, left=None, right=None)
    set_cell_vertical_align(logo_cell, "center")
    clear_cell_paras(logo_cell)
    lp = logo_cell.add_paragraph()
    no_space(lp)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        lp.add_run().add_picture(LOGO_PATH, width=Inches(1.25))
    else:
        add_run(lp, "CODESINO", bold=True, size_pt=14, color=WHITE_COLOR)

    # ── Title cell: slightly darker shade of the same blue ───────────────────
    set_cell_width(title_cell,    7160)
    set_cell_shading(title_cell,  "132F5A")
    set_cell_margins(title_cell,  top=240, bottom=240, left=260, right=300)
    set_cell_borders(title_cell,  top=None, bottom=None, left=None, right=None)
    set_cell_vertical_align(title_cell, "center")
    clear_cell_paras(title_cell)

    tp1 = title_cell.add_paragraph()
    add_paragraph_spacing(tp1, before=0, after=40)
    tp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(tp1, "CODESINO SOFTWARE DEVELOPMENT SERVICES",
            bold=True, size_pt=14, color=WHITE_COLOR)

    tp2 = title_cell.add_paragraph()
    add_paragraph_spacing(tp2, before=0, after=40)
    tp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(tp2, subtitle, size_pt=12, color=hex_to_rgb("C8DCFA"))

    tp3 = title_cell.add_paragraph()
    no_space(tp3)
    tp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(tp3, "www.codesinodev.com",
            italic=True, size_pt=9, color=hex_to_rgb("92B8E8"))

    spacer(doc, 60, 30)

    # ── Ref / Date bar ───────────────────────────────────────────────────────
    ref_table = doc.add_table(rows=1, cols=2)
    set_table_width(ref_table,    9360)
    set_column_widths(ref_table, [4680, 4680])
    set_table_no_borders(ref_table)

    for idx, (label, value, align) in enumerate([
        ("Agreement Reference",    agreement_ref,  WD_ALIGN_PARAGRAPH.LEFT),
        ("Date & Time Generated",  agreement_date, WD_ALIGN_PARAGRAPH.RIGHT),
    ]):
        cell = ref_table.cell(0, idx)
        set_cell_shading(cell,  HEX_LIGHT_BLUE)
        set_cell_margins(cell,  top=90, bottom=90, left=160, right=160)
        set_cell_borders(cell,
                         top    = {"val": "single", "sz": 4,  "color": HEX_MID_BLUE},
                         bottom = {"val": "single", "sz": 4,  "color": HEX_MID_BLUE},
                         left   = None, right = None)
        clear_cell_paras(cell)
        p = cell.add_paragraph()
        no_space(p)
        p.alignment = align
        add_run(p, f"{label}:  ", bold=True, size_pt=9.5, color=BRAND_BLUE)
        add_run(p, value,         bold=False, size_pt=9.5, color=DARK_GRAY)

    spacer(doc, 160, 60)


# ─────────────────────────────────────────────────────────────────────────────
# SHARED: SIGNATURE BLOCK  (employment agreements)
# ─────────────────────────────────────────────────────────────────────────────

def _build_signatures(doc, data):
    """
    Two-column signature table: Service Provider (left) | Employee (right).
    Provider cell includes stamp if the image exists.
    """
    heading1(doc, "AGREEMENT & SIGNATURES")
    body_text(doc,
              "By signing below, both parties confirm that they have read, understood, and "
              "agree to all the terms and conditions set forth in this Employment Agreement. "
              "This document constitutes a legally binding agreement between Codesino Software "
              "Development Services and the employee named herein.")
    spacer(doc, 80, 60)

    sig_table = doc.add_table(rows=2, cols=3)
    set_table_width(sig_table,    9360)
    set_column_widths(sig_table, [4400, 560, 4400])
    set_table_no_borders(sig_table)

    for col_idx, label, align in [
        (0, "CODESINO SOFTWARE DEVELOPMENT SERVICES", WD_ALIGN_PARAGRAPH.CENTER),
        (2, "EMPLOYEE",                               WD_ALIGN_PARAGRAPH.CENTER),
    ]:
        cell = sig_table.cell(0, col_idx)
        set_cell_shading(cell,  HEX_BRAND_BLUE)
        set_cell_margins(cell,  top=100, bottom=100, left=150, right=150)
        set_cell_borders(cell,  top=None, bottom=None, left=None, right=None)
        clear_cell_paras(cell)
        p = cell.add_paragraph()
        no_space(p)
        p.alignment = align
        add_run(p, label, bold=True, size_pt=10, color=WHITE_COLOR)

    gap_hdr = sig_table.cell(0, 1)
    set_cell_borders(gap_hdr, top=None, bottom=None, left=None, right=None)
    clear_cell_paras(gap_hdr)
    gap_hdr.add_paragraph()

    thin = {"val": "single", "sz": 2, "color": HEX_RULE}

    # Provider cell
    prov_cell = sig_table.cell(1, 0)
    set_cell_borders(prov_cell, top=None, bottom=thin, left=thin, right=thin)
    set_cell_margins(prov_cell, top=120, bottom=140, left=150, right=150)
    clear_cell_paras(prov_cell)
    pp1 = prov_cell.add_paragraph()
    add_paragraph_spacing(pp1, before=0, after=50)
    add_run(pp1, "Codesino Software Development Services", size_pt=10, color=DARK_GRAY)
    pp2 = prov_cell.add_paragraph()
    add_paragraph_spacing(pp2, before=0, after=80)
    add_run(pp2, "Authorised Signatory / Management", size_pt=10, color=MID_GRAY)
    sig_line(prov_cell, "Digital Signature / Stamp:")
    if os.path.exists(STAMP_PATH):
        stamp_p = prov_cell.add_paragraph()
        add_paragraph_spacing(stamp_p, before=20, after=20)
        stamp_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        stamp_p.add_run().add_picture(STAMP_PATH, width=Inches(1.3))
    pds = data.get("providerSignDate", "")
    pp4 = prov_cell.add_paragraph()
    no_space(pp4)
    add_run(pp4, f"Date: {format_date(pds) if pds else '___ / ___ / ______'}",
            italic=True, size_pt=10, color=hex_to_rgb("999999"))

    gap_body = sig_table.cell(1, 1)
    set_cell_borders(gap_body, top=None, bottom=None, left=None, right=None)
    clear_cell_paras(gap_body)
    gap_body.add_paragraph()

    # Employee cell
    emp_cell = sig_table.cell(1, 2)
    set_cell_borders(emp_cell, top=None, bottom=thin, left=thin, right=thin)
    set_cell_margins(emp_cell, top=120, bottom=140, left=150, right=150)
    clear_cell_paras(emp_cell)
    en = data.get("employeeName", "")
    ep1 = emp_cell.add_paragraph()
    add_paragraph_spacing(ep1, before=0, after=50)
    add_run(ep1, en if en else "____________________________",
            italic=not bool(en), size_pt=10,
            color=DARK_GRAY if en else hex_to_rgb("999999"))
    ep2 = emp_cell.add_paragraph()
    add_paragraph_spacing(ep2, before=0, after=80)
    add_run(ep2, "Employee — Full Name", size_pt=10, color=MID_GRAY)
    sig_line(emp_cell, "Signature:")
    blank_p = emp_cell.add_paragraph()
    add_paragraph_spacing(blank_p, before=20, after=20)
    no_space(blank_p)
    eds = data.get("employeeSignDate", "")
    ep4 = emp_cell.add_paragraph()
    no_space(ep4)
    add_run(ep4, f"Date: {format_date(eds) if eds else '___ / ___ / ______'}",
            italic=True, size_pt=10, color=hex_to_rgb("999999"))

    spacer(doc, 100, 60)


# ─────────────────────────────────────────────────────────────────────────────
# SHARED: DOCUMENT FOOTER  (employment agreements)
# ─────────────────────────────────────────────────────────────────────────────

def _build_footer(doc, doc_type_label):
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(fp)
    add_run(fp,
            "This agreement was prepared by Codesino Software Development Services  |  "
            "www.codesinodev.com  |  contact@codesinodev.com",
            italic=True, size_pt=9, color=hex_to_rgb("9E9E9E"))

    section = doc.sections[0]
    footer  = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_para,
            f"Codesino Software Development Services  |  {doc_type_label}  |  CONFIDENTIAL",
            size_pt=8, color=hex_to_rgb("999999"))


# ─────────────────────────────────────────────────────────────────────────────
# SHARED: EMPLOYMENT DETAILS SECTION  (common to all three employment types)
# ─────────────────────────────────────────────────────────────────────────────

def _build_employment_details(doc, data, position_label, positions_list,
                              allow_multi_position=False):
    """
    Renders Section 2 — Employment Details.
    position_label: "Position / Role"
    positions_list: the relevant positions catalogue for checkboxes
    allow_multi_position: if True, renders each position as its own checkbox row
    """
    heading1(doc, "2.  EMPLOYMENT DETAILS")
    spacer(doc, 40, 40)

    emp_type = data.get("employmentType", "")

    def emp_details_fn(cell):
        lv_para(cell, "Employee Full Name",   data.get("employeeName",     ""))
        lv_para(cell, "Personal Email",       data.get("employeeEmail",    ""))
        lv_para(cell, "Phone / WhatsApp",     data.get("employeePhone",    ""))
        lv_para(cell, "Residential Address",  data.get("employeeAddress",  ""))
        lv_para(cell, "Date of Hire",         format_date(data.get("dateHired", "")) if data.get("dateHired") else "")
        lv_para(cell, "Interviewed / Onboarded By", data.get("interviewedBy", ""))
        lv_para(cell, "Reporting To",         data.get("reportingTo",      ""))
        lv_para(cell, "Work Location / Mode", data.get("workMode",         ""))
        p = cell.add_paragraph()
        add_paragraph_spacing(p, before=0, after=56)
        add_run(p, "Employment Type:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        type_str = "    ".join(f"{cb(emp_type == t)}  {t}" for t in EMPLOYMENT_TYPES)
        add_run(p, type_str, size_pt=11, color=DARK_GRAY)
        if data.get("employmentTypeOther"):
            lv_para(cell, "Other (specified)", data.get("employmentTypeOther", ""))

    section_box(doc, emp_details_fn, HEX_WHITE)
    spacer(doc, 80, 60)

    heading2(doc, "2.1  Position / Role Assigned")

    if allow_multi_position:
        selected_positions = data.get("positions", [])
        if isinstance(selected_positions, str):
            selected_positions = [selected_positions]
        other_pos = data.get("positionOther", "")

        def pos_fn(cell):
            hp = cell.add_paragraph()
            add_paragraph_spacing(hp, before=0, after=40)
            add_run(hp, "Selected Position(s): ", bold=True, size_pt=11, color=BRAND_BLUE)
            add_run(hp, "(employee may hold more than one role)",
                    italic=True, size_pt=10, color=MID_GRAY)
            for pos in positions_list:
                pp = cell.add_paragraph()
                add_paragraph_spacing(pp, before=10, after=10)
                is_sel = pos in selected_positions
                add_run(pp, f"{cb(is_sel)}  {pos}", size_pt=11,
                        color=DARK_GRAY if not is_sel else BRAND_BLUE,
                        bold=is_sel)
            if "Other (Specified)" in selected_positions and other_pos:
                lv_para(cell, "Other (specified)", other_pos)
            lv_para(cell, "Team / Department", data.get("department", ""))

        section_box(doc, pos_fn, HEX_WHITE)
    else:
        sel_pos   = data.get("position", "")
        other_pos = data.get("positionOther", "")

        def pos_fn(cell):
            hp = cell.add_paragraph()
            add_paragraph_spacing(hp, before=0, after=40)
            add_run(hp, "Position: ", bold=True, size_pt=11, color=BRAND_BLUE)
            for pos in positions_list:
                pp = cell.add_paragraph()
                add_paragraph_spacing(pp, before=10, after=10)
                is_sel = (sel_pos == pos)
                add_run(pp, f"{cb(is_sel)}  {pos}", size_pt=11,
                        color=DARK_GRAY if not is_sel else BRAND_BLUE,
                        bold=is_sel)
            if sel_pos == "Other (Specified)" and other_pos:
                lv_para(cell, "Other (specified)", other_pos)
            lv_para(cell, "Department / Division", data.get("department", ""))

        section_box(doc, pos_fn, HEX_WHITE)

    divider(doc)


# ─────────────────────────────────────────────────────────────────────────────
# SHARED: COMPENSATION SECTION
# ─────────────────────────────────────────────────────────────────────────────

def _build_compensation(doc, data, currency_sym):
    heading1(doc, "3.  COMPENSATION & BENEFITS")
    spacer(doc, 40, 40)

    sal_type = data.get("salaryType", "")

    def comp_fn(cell):
        p = cell.add_paragraph()
        add_paragraph_spacing(p, before=0, after=56)
        add_run(p, "Salary / Compensation Structure:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        sal_str = "    ".join(f"{cb(sal_type == s)}  {s}" for s in SALARY_TYPES)
        add_run(p, sal_str, size_pt=11, color=DARK_GRAY)

        if sal_type == "Monthly Fixed Salary":
            lv_para(cell, "Monthly Gross Salary",
                    f"{currency_sym}{data.get('salaryAmount', '')} / month")
        elif sal_type == "Percentage Per Contract":
            lv_para(cell, "Commission Rate",
                    f"{data.get('salaryPercent', '')}% per completed contract/project")
        elif sal_type == "Hourly Rate":
            lv_para(cell, "Hourly Rate",
                    f"{currency_sym}{data.get('salaryAmount', '')} / hour")
        elif sal_type == "Commission-Based":
            lv_para(cell, "Commission Structure", data.get("commissionStructure", ""))
        elif sal_type == "Retainer Fee":
            lv_para(cell, "Monthly Retainer",
                    f"{currency_sym}{data.get('salaryAmount', '')} / month")

        lv_para(cell, "Payment Frequency",    data.get("paymentFrequency",  ""))
        lv_para(cell, "Payment Method",       data.get("paymentMethod",     ""))
        lv_para(cell, "Probation Period",     data.get("probationPeriod",   ""))
        lv_para(cell, "Benefits / Allowances", data.get("benefits",         ""))
        lv_para(cell, "Working Hours / Days", data.get("workingHours",      ""))
        lv_para(cell, "Additional Notes on Compensation", data.get("compensationNotes", ""))

    section_box(doc, comp_fn, HEX_PALE_BLUE, accent_color=HEX_MID_BLUE)
    divider(doc)


# ─────────────────────────────────────────────────────────────────────────────
# SHARED: STANDARD LEGAL CLAUSES
# ─────────────────────────────────────────────────────────────────────────────

def _build_standard_clauses(doc):
    heading1(doc, "INTELLECTUAL PROPERTY & OWNERSHIP")
    for item in [
        "Any work, code, content, design, strategy, or deliverable produced by the employee "
        "in the course of their employment with Codesino Software Development Services is the "
        "sole intellectual property of Codesino Software Development Services.",
        "The employee shall not reproduce, distribute, or commercialise any work product "
        "created during employment without express written consent from the company.",
        "All intellectual property rights vest entirely in the company from the moment of "
        "creation, regardless of where or when the work was performed.",
        "Any pre-existing intellectual property owned by the employee and used during the "
        "course of work must be declared in writing prior to use.",
        "These IP obligations survive the termination of this agreement indefinitely.",
    ]:
        add_numbered(doc, item)
    divider(doc)

    heading1(doc, "TERMINATION OF EMPLOYMENT")
    for item in [
        "Either party may terminate this agreement with a minimum of 14 (fourteen) "
        "calendar days' written notice, unless otherwise stipulated by applicable law.",
        "The company reserves the right to terminate employment immediately and without "
        "notice in cases of gross misconduct, breach of confidentiality, fraud, or "
        "actions that cause reputational or financial harm to the company or its clients.",
        "Upon termination, the employee shall immediately return all company property, "
        "devices, credentials, access tokens, and confidential materials.",
        "Outstanding compensation for work already completed up to the termination date "
        "shall be settled within 30 (thirty) business days of the effective termination date.",
        "Post-termination, the employee remains bound by the confidentiality obligations "
        "outlined in this agreement for a period of 2 (two) years.",
    ]:
        add_numbered(doc, item)
    divider(doc)

    heading1(doc, "DISPUTE RESOLUTION")
    body_text(doc,
              "In the event of any dispute arising under this agreement, both parties agree "
              "to first attempt resolution through good-faith negotiation within 14 (fourteen) "
              "calendar days of the dispute being raised in writing. If an amicable resolution "
              "cannot be achieved, both parties agree to submit the matter to mediation before "
              "resorting to formal legal proceedings. This agreement shall be governed by and "
              "construed in accordance with the laws of the Federal Republic of Nigeria.")
    divider(doc)


# ─────────────────────────────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════
# GENERATOR 1 — CLIENT PROJECT AGREEMENT
# ═══════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────

def generate_agreement_buffer(data: dict) -> bytes:
    doc = _setup_document()

    agreement_ref  = generate_ref("AGR")
    agreement_date = format_datetime()
    currency       = data.get("currency", "ngn")
    currency_sym   = "$" if currency == "usd" else "₦"

    # ═════════════════════════════════════════════════════════════════════════
    # COVER HEADER  –  logo + title block + ref/date bar
    # ═════════════════════════════════════════════════════════════════════════

    # ── Top bar: logo (left) | title text (right) ────────────────────────────
    hdr_table = doc.add_table(rows=1, cols=2)
    set_table_width(hdr_table,    9360)
    set_column_widths(hdr_table, [2200, 7160])
    set_table_no_borders(hdr_table)

    logo_cell  = hdr_table.cell(0, 0)
    title_cell = hdr_table.cell(0, 1)

    # Logo cell — white background, logo image centred
    set_cell_width(logo_cell,    2200)
    set_cell_shading(logo_cell,  HEX_WHITE)
    set_cell_margins(logo_cell,  top=100, bottom=100, left=100, right=100)
    set_cell_borders(logo_cell,  top=None, bottom=None, left=None, right=None)
    set_cell_vertical_align(logo_cell, "center")
    clear_cell_paras(logo_cell)
    lp = logo_cell.add_paragraph()
    no_space(lp)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        lp.add_run().add_picture(LOGO_PATH, width=Inches(1.35))
    else:
        add_run(lp, "CODESINO", bold=True, size_pt=14, color=BRAND_BLUE)

    # Title cell — deep blue background
    set_cell_width(title_cell,    7160)
    set_cell_shading(title_cell,  HEX_BRAND_BLUE)
    set_cell_margins(title_cell,  top=240, bottom=240, left=260, right=300)
    set_cell_borders(title_cell,  top=None, bottom=None, left=None, right=None)
    set_cell_vertical_align(title_cell, "center")
    clear_cell_paras(title_cell)

    tp1 = title_cell.add_paragraph()
    add_paragraph_spacing(tp1, before=0, after=40)
    tp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(tp1, "CODESINO SOFTWARE DEVELOPMENT SERVICES",
            bold=True, size_pt=15, color=WHITE_COLOR)

    tp2 = title_cell.add_paragraph()
    add_paragraph_spacing(tp2, before=0, after=40)
    tp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(tp2, "Client Project Agreement",
            size_pt=12, color=hex_to_rgb("C8DCFA"))

    tp3 = title_cell.add_paragraph()
    no_space(tp3)
    tp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(tp3, "www.codesinodev.com",
            italic=True, size_pt=9, color=hex_to_rgb("92B8E8"))

    spacer(doc, 60, 30)

    # ── Ref / Date bar ───────────────────────────────────────────────────────
    ref_table = doc.add_table(rows=1, cols=2)
    set_table_width(ref_table,    9360)
    set_column_widths(ref_table, [4680, 4680])
    set_table_no_borders(ref_table)

    for idx, (label, value, align) in enumerate([
        ("Agreement Reference", agreement_ref,  WD_ALIGN_PARAGRAPH.LEFT),
        ("Date & Time Generated", agreement_date, WD_ALIGN_PARAGRAPH.RIGHT),
    ]):
        cell = ref_table.cell(0, idx)
        set_cell_shading(cell,  HEX_LIGHT_BLUE)
        set_cell_margins(cell,  top=90, bottom=90, left=160, right=160)
        set_cell_borders(cell,
                         top    = {"val": "single", "sz": 4,  "color": HEX_MID_BLUE},
                         bottom = {"val": "single", "sz": 4,  "color": HEX_MID_BLUE},
                         left   = None, right = None)
        clear_cell_paras(cell)
        p = cell.add_paragraph()
        no_space(p)
        p.alignment = align
        add_run(p, f"{label}:  ", bold=True, size_pt=9.5, color=BRAND_BLUE)
        add_run(p, value,         bold=False, size_pt=9.5, color=DARK_GRAY)

    spacer(doc, 160, 60)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 1 — PARTIES
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "1.  PARTIES TO THIS AGREEMENT")
    spacer(doc, 40, 40)
    heading2(doc, "1.1  Service Provider")

    def provider_fn(cell):
        lv_para(cell, "Company Name",       "Codesino Software Development Services")
        lv_para(cell, "Email",              "contact@codesinodev.com")
        lv_para(cell, "Phone / WhatsApp",   "+2349036206457")
        lv_para(cell, "Website",            "https://www.codesinodev.com")

    section_box(doc, provider_fn, HEX_WHITE)
    spacer(doc, 80, 60)
    heading2(doc, "1.2  Client")

    def client_fn(cell):
        lv_para(cell, "Client Full Name",         data.get("clientName",    ""))
        lv_para(cell, "Company / Organisation",   data.get("clientCompany", ""))
        lv_para(cell, "Email Address",            data.get("clientEmail",   ""))
        lv_para(cell, "Phone / WhatsApp",         data.get("clientPhone",   ""))

    section_box(doc, client_fn, HEX_WHITE)
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 2 — PROJECT OVERVIEW
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "2.  PROJECT OVERVIEW")
    spacer(doc, 40, 40)

    proj_types    = ["Website", "Web Application", "E-Commerce", "Web Portal", "Other"]
    sel_type      = data.get("projectType", "")
    type_cbs      = "      ".join(f"{cb(sel_type == t)}  {t}" for t in proj_types)

    def overview_fn(cell):
        lv_para(cell, "Project Name", data.get("projectName", ""))
        p = cell.add_paragraph()
        add_paragraph_spacing(p, before=0, after=56)
        add_run(p, "Project Type:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        add_run(p, type_cbs, size_pt=11, color=DARK_GRAY)
        if sel_type == "Other" and data.get("projectTypeOther"):
            lv_para(cell, "Other (specified)", data.get("projectTypeOther", ""))
        dp = cell.add_paragraph()
        add_paragraph_spacing(dp, before=0, after=30)
        add_run(dp, "Project Description:", bold=True, size_pt=11, color=BRAND_BLUE)
        desc = data.get("projectDescription", "")
        ddp  = cell.add_paragraph()
        add_paragraph_spacing(ddp, before=0, after=56, line=276)
        add_run(ddp, desc if desc else "Not Provided",
                italic=not bool(desc), size_pt=11,
                color=DARK_GRAY if desc else PLACEHOLDER)
        lv_para(cell, "Target Audience / End Users", data.get("targetAudience", ""))

    section_box(doc, overview_fn, HEX_WHITE)
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 3 — SCOPE OF WORK
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "3.  SCOPE OF WORK")
    italic_note(doc, "All features listed below have been agreed upon by both parties. Any feature not listed is considered out of scope and may attract additional charges.")
    spacer(doc, 40, 40)

    heading2(doc, "3.1  Core / Main Features")
    core = [f for f in (data.get("coreFeatures") or []) if f.get("feature")]
    if core:
        data_table(doc,
                   ["#", "Feature / Deliverable", "Description"],
                   [[str(i + 1), f["feature"], f.get("description", "—")] for i, f in enumerate(core)],
                   [480, 3840, 5040])
    else:
        italic_note(doc, "No core features specified.")

    spacer(doc, 120, 60)
    heading2(doc, "3.2  Extra / Add-on Features")
    add_ons = [f for f in (data.get("addOnFeatures") or []) if f.get("feature")]
    if add_ons and data.get("hasAddOns"):
        data_table(doc,
                   ["#", "Extra Feature", "Notes / Condition"],
                   [[str(i + 1), f["feature"], f.get("notes", "—")] for i, f in enumerate(add_ons)],
                   [480, 3840, 5040])
    else:
        italic_note(doc, "No add-on features specified for this agreement.")

    spacer(doc, 120, 60)
    heading2(doc, "3.3  Explicitly Out of Scope")
    italic_note(doc, "The following items are NOT included in this agreement. Any additional requests will be quoted separately.")
    oos = [x for x in (data.get("outOfScope") or []) if x]
    if oos:
        for item in oos:
            add_numbered(doc, item)
    else:
        italic_note(doc, "No explicit out-of-scope items listed.")

    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 4 — DESIGN SPECIFICATIONS
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "4.  UI / DESIGN SPECIFICATIONS")
    spacer(doc, 40, 40)

    des_styles  = ["Minimal", "Corporate", "Bold/Vibrant", "Elegant", "Playful", "Other"]
    sel_style   = data.get("designStyle", "")
    style_cbs   = "    ".join(f"{cb(sel_style == s)}  {s}" for s in des_styles)
    logo_prov   = data.get("logoProvided", "")
    brand_assets= data.get("brandAssetsProvided", "")

    def design_fn(cell):
        lv_para(cell, "Primary Brand Colour",    data.get("primaryColor",   ""))
        lv_para(cell, "Secondary Colour",         data.get("secondaryColor", ""))
        lv_para(cell, "Accent Colour",            data.get("accentColor",    ""))
        lv_para(cell, "Preferred Font(s)",        data.get("preferredFonts", ""))
        p = cell.add_paragraph()
        add_paragraph_spacing(p, before=0, after=56)
        add_run(p, "Design Style / Mood:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        add_run(p, style_cbs, size_pt=11, color=DARK_GRAY)
        if sel_style == "Other" and data.get("designStyleOther"):
            lv_para(cell, "Other (specified)", data.get("designStyleOther", ""))
        lv_para(cell, "Inspiration Websites", data.get("inspirationSites", ""))
        lp = cell.add_paragraph()
        add_paragraph_spacing(lp, before=0, after=56)
        add_run(lp, "Logo Provided by Client:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        add_run(lp, f"{cb(logo_prov == 'yes')}  Yes      {cb(logo_prov == 'no')}  No (Codesino will create basic logo — quoted separately)",
                size_pt=11, color=DARK_GRAY)
        bp = cell.add_paragraph()
        add_paragraph_spacing(bp, before=0, after=56)
        add_run(bp, "Brand Assets / Style Guide Provided:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        add_run(bp, f"{cb(brand_assets == 'yes')}  Yes      {cb(brand_assets == 'no')}  No",
                size_pt=11, color=DARK_GRAY)
        rev = data.get("revisionRounds", "")
        lv_para(cell, "Design Revision Rounds",
                f"{rev} rounds (additional revisions billed separately)" if rev else "")

    section_box(doc, design_fn, HEX_WHITE)
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 5 — PAGES & SCREENS
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "5.  PAGES & SCREENS INCLUDED")
    italic_note(doc, "All pages listed below are included in this agreement. Any page not listed is out of scope.")
    spacer(doc, 40, 40)

    default_pages = [
        {"name": "Home / Landing Page",       "notes": "Default"},
        {"name": "About Page",                "notes": "Default"},
        {"name": "Services / Products Page",  "notes": "Default"},
        {"name": "Contact Page",              "notes": "Default"},
    ]
    extra_pages = [p for p in (data.get("extraPages") or []) if p.get("name")]
    all_pages   = default_pages + [{"name": p["name"], "notes": p.get("notes", "Additional")}
                                    for p in extra_pages]
    data_table(doc,
               ["#", "Page / Screen Name", "Notes"],
               [[str(i + 1), p["name"], p["notes"]] for i, p in enumerate(all_pages)],
               [480, 4440, 4440])
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 6 — TECHNICAL SPECIFICATIONS
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "6.  TECHNICAL SPECIFICATIONS")
    spacer(doc, 40, 40)

    dom_by   = data.get("domainProvidedBy", "")
    ssl_stat = data.get("sslCertificate",   "")
    mob_resp = data.get("mobileResponsive", "")

    def tech_fn(cell):
        lv_para(cell, "Frontend Technology",  data.get("frontendTech",    ""))
        lv_para(cell, "Backend Technology",   data.get("backendTech",     ""))
        lv_para(cell, "Database",             data.get("database",        ""))
        lv_para(cell, "Hosting Platform",     data.get("hostingPlatform", ""))
        lv_para(cell, "Domain Name",          data.get("domainName",      ""))
        dp = cell.add_paragraph()
        add_paragraph_spacing(dp, before=0, after=56)
        add_run(dp, "Domain Provided By:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        add_run(dp, f"{cb(dom_by == 'client')}  Client      {cb(dom_by == 'codesino')}  Codesino Development",
                size_pt=11, color=DARK_GRAY)
        sp = cell.add_paragraph()
        add_paragraph_spacing(sp, before=0, after=56)
        add_run(sp, "SSL Certificate:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        add_run(sp, f"{cb(ssl_stat == 'included')}  Included      {cb(ssl_stat == 'not-included')}  Not Included      {cb(ssl_stat == 'client')}  Client to Arrange",
                size_pt=11, color=DARK_GRAY)
        mp = cell.add_paragraph()
        no_space(mp)
        add_run(mp, "Mobile Responsive:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        add_run(mp, f"{cb(mob_resp == 'yes')}  Yes      {cb(mob_resp == 'no')}  No",
                size_pt=11, color=DARK_GRAY)

    section_box(doc, tech_fn, HEX_WHITE)
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 7 — DEVELOPER BRIEFS
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "7.  INTERNAL DEVELOPER BRIEFS")
    italic_note(doc, "⚠  INTERNAL USE ONLY — This section is for the Codesino development team and is not intended for the client.")
    spacer(doc, 60, 40)

    heading2(doc, "7.1  Frontend Developer Notes")

    def fe_fn(cell):
        lv_para(cell, "Assigned Developer", data.get("frontendDev", ""))
        hp = cell.add_paragraph()
        add_paragraph_spacing(hp, before=40, after=40)
        add_run(hp, "Default Responsibilities:", bold=True, size_pt=11, color=BRAND_BLUE)
        for item in [
            "Build all pages per the UI specs in Section 4",
            "Ensure full mobile responsiveness across all standard breakpoints",
            "Integrate with backend APIs provided by the backend developer",
            "Cross-browser compatibility (Chrome, Firefox, Safari, Edge)",
            "Implement SEO-friendly markup and performance best practices",
        ]:
            add_bullet(cell, item)
        np = cell.add_paragraph()
        add_paragraph_spacing(np, before=60, after=30)
        add_run(np, "Special Notes / Extra Instructions:", bold=True, size_pt=11, color=BRAND_BLUE)
        fn  = data.get("frontendNotes", "")
        fnp = cell.add_paragraph()
        no_space(fnp)
        add_run(fnp, fn if fn else "None", italic=not bool(fn), size_pt=11,
                color=DARK_GRAY if fn else PLACEHOLDER)

    section_box(doc, fe_fn, HEX_LIGHT_BLUE, accent_color=HEX_MID_BLUE)
    spacer(doc, 80, 60)

    heading2(doc, "7.2  Backend Developer Notes")

    def be_fn(cell):
        lv_para(cell, "Assigned Developer", data.get("backendDev", ""))
        hp = cell.add_paragraph()
        add_paragraph_spacing(hp, before=40, after=40)
        add_run(hp, "Default Responsibilities:", bold=True, size_pt=11, color=BRAND_BLUE)
        for item in [
            "Set up and configure server, database, and hosting environment",
            "Build and document all API endpoints required by the frontend",
            "Handle authentication, data validation, and security hardening",
            "Database schema design and migration management",
            "Implement error handling, logging, and monitoring",
        ]:
            add_bullet(cell, item)
        np = cell.add_paragraph()
        add_paragraph_spacing(np, before=60, after=30)
        add_run(np, "Special Notes / Extra Instructions:", bold=True, size_pt=11, color=BRAND_BLUE)
        bn  = data.get("backendNotes", "")
        bnp = cell.add_paragraph()
        no_space(bnp)
        add_run(bnp, bn if bn else "None", italic=not bool(bn), size_pt=11,
                color=DARK_GRAY if bn else PLACEHOLDER)

    section_box(doc, be_fn, HEX_LIGHT_BLUE, accent_color=HEX_MID_BLUE)
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 8 — TIMELINE
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "8.  PROJECT TIMELINE")
    italic_note(doc, "Timelines are subject to client feedback response times and payment completion.")
    spacer(doc, 40, 40)

    start_label = (format_date(data.get("startDate")) + " (subject to payment confirmation)") \
                  if data.get("startDate") else "Not agreed"
    end_label   = format_date(data.get("endDate")) if data.get("endDate") else "Not agreed"

    def timeline_fn(cell):
        lv_para(cell, "Estimated Start Date",       start_label)
        lv_para(cell, "Estimated Completion Date",  end_label)
        lv_para(cell, "Total Development Duration", data.get("duration", ""))

    section_box(doc, timeline_fn, HEX_WHITE)
    spacer(doc, 80, 60)

    heading2(doc, "8.1  Project Milestones")
    milestones = [
        ("Project Kickoff (payment confirmed)",   data.get("m1Date", "")),
        ("Design Mockups / Wireframes",            data.get("m2Date", "")),
        ("Client Design Approval",                 data.get("m3Date", "")),
        ("Frontend Development Complete",          data.get("m4Date", "")),
        ("Backend Development Complete",           data.get("m5Date", "")),
        ("Integration & Testing",                  data.get("m6Date", "")),
        ("Client Review & Feedback",               data.get("m7Date", "")),
        ("Final Revisions",                        data.get("m8Date", "")),
        ("Deployment / Go Live",                   data.get("m9Date", "")),
    ]
    data_table(doc,
               ["Milestone", "Expected Completion", "Status"],
               [[m, format_date(d) if d else "TBC", "☐  Pending"] for m, d in milestones],
               [4000, 3000, 2360])

    spacer(doc, 80, 60)
    heading2(doc, "8.2  Timeline Policy")
    body_text(doc,
              "Development does not begin until the required upfront payment is confirmed. "
              "The timeline starts from the date payment is received, not the date this "
              "agreement is signed. Codesino is not responsible for delays caused by late "
              "client feedback, delayed provision of content or assets, change of scope "
              "after sign-off, or unavailability of third-party services.")
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 9 — PRICING & PAYMENT
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "9.  PRICING & PAYMENT TERMS")
    italic_note(doc, "No development work begins until the required payment is received and confirmed in writing.")
    spacer(doc, 60, 40)

    total_amount  = data.get("totalAmount", "")
    payment_plan  = data.get("paymentPlan", "")

    def pricing_fn(cell):
        tp = cell.add_paragraph()
        add_paragraph_spacing(tp, before=0, after=56)
        add_run(tp, "Total Project Amount:  ", bold=True, size_pt=14, color=BRAND_BLUE)
        add_run(tp, f"{currency_sym}{total_amount}", bold=True, size_pt=14, color=DARK_GRAY)
        lv_para(cell, "Amount in Words", data.get("amountInWords", ""))

    section_box(doc, pricing_fn, HEX_PALE_BLUE, accent_color=HEX_MID_BLUE)
    spacer(doc, 80, 60)
    heading2(doc, "9.1  Payment Terms")

    # Option A
    def opt_a_fn(cell):
        is_sel = payment_plan == "A"
        p = cell.add_paragraph()
        add_paragraph_spacing(p, before=0, after=56)
        add_run(p, f"{cb(is_sel)}  OPTION A: Full Payment Before Kickoff",
                bold=True, size_pt=12,
                color=GREEN_ACCENT if is_sel else hex_to_rgb("888888"))
        d = cell.add_paragraph()
        add_paragraph_spacing(d, before=0, after=40)
        add_run(d, "The client pays 100% of the total project amount before any development begins.",
                size_pt=11, color=DARK_GRAY)
        if is_sel:
            lv_para(cell, "Amount Due Upfront", f"{currency_sym}{total_amount}  (100%)")
            sp = cell.add_paragraph()
            add_paragraph_spacing(sp, before=0, after=40)
            st = data.get("optionAStatus", "")
            add_run(sp, "Payment Status:  ", bold=True, size_pt=11, color=DARK_GRAY)
            add_run(sp, f"{cb(st == 'received')}  Received & Confirmed        {cb(st == 'pending')}  Pending",
                    size_pt=11, color=DARK_GRAY)
            if st == "received":
                lv_para(cell, "Confirmation Date",
                        format_date(data.get("optionADate", "")) if data.get("optionADate") else "")
                lv_para(cell, "Confirmed Amount", f"{currency_sym}{data.get('optionAAmount', '')}")
            np = cell.add_paragraph()
            no_space(np)
            add_run(np, "Note: Development begins only after this payment is fully confirmed.",
                    italic=True, size_pt=10, color=MID_GRAY)

    section_box(doc, opt_a_fn, HEX_F0F8FF if payment_plan == "A" else HEX_FAFAFA,
                accent_color=HEX_GREEN if payment_plan == "A" else HEX_RULE)
    spacer(doc, 80, 60)

    # Option B
    def opt_b_fn(cell):
        is_sel = payment_plan == "B"
        p = cell.add_paragraph()
        add_paragraph_spacing(p, before=0, after=56)
        add_run(p, f"{cb(is_sel)}  OPTION B: 60% Upfront / 40% on Completion",
                bold=True, size_pt=12,
                color=BRAND_BLUE if is_sel else hex_to_rgb("888888"))
        d = cell.add_paragraph()
        add_paragraph_spacing(d, before=0, after=40)
        add_run(d, "The client pays 60% before development begins. The remaining 40% is due upon completion, before deployment.",
                size_pt=11, color=DARK_GRAY)
        if is_sel:
            h1 = cell.add_paragraph()
            add_paragraph_spacing(h1, before=40, after=30)
            add_run(h1, "First Payment — 60%  (Due Before Kickoff)", bold=True, size_pt=11, color=BRAND_BLUE)
            lv_para(cell, "Amount", f"{currency_sym}{data.get('optionB1Amount', '')}")
            s1p = cell.add_paragraph()
            add_paragraph_spacing(s1p, before=0, after=40)
            s1 = data.get("optionB1Status", "")
            add_run(s1p, "Status:  ", bold=True, size_pt=11, color=DARK_GRAY)
            add_run(s1p, f"{cb(s1 == 'received')}  Received & Confirmed        {cb(s1 == 'pending')}  Pending",
                    size_pt=11, color=DARK_GRAY)
            if s1 == "received":
                lv_para(cell, "Confirmation Date",
                        format_date(data.get("optionB1Date", "")) if data.get("optionB1Date") else "")

            h2 = cell.add_paragraph()
            add_paragraph_spacing(h2, before=60, after=30)
            add_run(h2, "Second Payment — 40%  (Due Before Deployment)", bold=True, size_pt=11, color=BRAND_BLUE)
            lv_para(cell, "Amount", f"{currency_sym}{data.get('optionB2Amount', '')}")
            s2p = cell.add_paragraph()
            add_paragraph_spacing(s2p, before=0, after=40)
            s2 = data.get("optionB2Status", "")
            add_run(s2p, "Status:  ", bold=True, size_pt=11, color=DARK_GRAY)
            add_run(s2p, f"{cb(s2 == 'received')}  Received & Confirmed        {cb(s2 == 'pending')}  Pending",
                    size_pt=11, color=DARK_GRAY)
            if s2 == "received":
                lv_para(cell, "Confirmation Date",
                        format_date(data.get("optionB2Date", "")) if data.get("optionB2Date") else "")

            np = cell.add_paragraph()
            no_space(np)
            add_run(np, "Note: The final 40% must be received before deployment. No project goes live on an outstanding balance.",
                    italic=True, size_pt=10, color=MID_GRAY)

    section_box(doc, opt_b_fn, HEX_LIGHT_BLUE if payment_plan == "B" else HEX_FAFAFA,
                accent_color=HEX_BRAND_BLUE if payment_plan == "B" else HEX_RULE)
    spacer(doc, 80, 60)

    heading2(doc, "9.2  Payment Account Details")

    def ngn_fn(cell):
        lv_para(cell, "Bank Name",       "KudaBank")
        lv_para(cell, "Account Name",    "Codesino Software Development Services")
        lv_para(cell, "Account Number",  "3003017268")
        lv_para(cell, "Currency",        "Nigerian Naira (₦)")

    def usd_fn(cell):
        lv_para(cell, "Bank Name",      "— (USD account details pending)")
        lv_para(cell, "Account Name",   "— (to be added)")
        lv_para(cell, "Account Number", "— (to be added)")
        lv_para(cell, "Currency",       "US Dollar ($)")

    section_box(doc, ngn_fn if currency == "ngn" else usd_fn, HEX_WHITE)
    spacer(doc, 80, 60)

    heading2(doc, "9.3  General Payment Policy")
    payment_policy = [
        "No development work commences until the required upfront payment has been received and confirmed in writing by Codesino Software Development Services.",
        "All payments made are strictly non-refundable once the development phase has commenced.",
        "If the client cancels the project after development has started, all payments made up to that point are forfeited.",
        "If Codesino fails to deliver the agreed scope within the agreed timeline (excluding client-caused delays), a partial or full refund may be negotiated at the discretion of both parties.",
        "All amounts are in the currency agreed above. Cross-currency payment requests are subject to conversion at the prevailing rate on the date of payment.",
        "Codesino reserves the right to suspend work on any project where payment is overdue by more than 7 (seven) business days.",
        "A receipt/confirmation will be issued for every payment received.",
    ]
    for item in payment_policy:
        add_numbered(doc, item)
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 10 — CONTENT & ASSETS
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "10.  CONTENT & ASSETS")
    italic_note(doc, "Codesino builds the structure and functionality. Content delays by the client will result in corresponding timeline delays.")
    spacer(doc, 40, 40)

    asset_data  = data.get("assets") or {}
    asset_items = [
        ("Website copy / written content",  asset_data.get("websiteCopy")),
        ("Product / service images",        asset_data.get("productImages")),
        ("Company logo (PNG / SVG)",        asset_data.get("companyLogo")),
        ("Video content",                   asset_data.get("videoContent")),
        ("Staff / team photos",             asset_data.get("teamPhotos")),
        ("Social media links",              asset_data.get("socialLinks")),
    ]
    extra_assets = [a for a in (data.get("extraAssets") or []) if a.get("name")]
    for a in extra_assets:
        asset_items.append((a["name"], a.get("providedBy", "")))

    def asset_label(by):
        if by == "client":   return "☑  Client   ☐  Codesino"
        if by == "codesino": return "☐  Client   ☑  Codesino"
        return "☐  Client   ☐  Codesino"

    data_table(doc,
               ["Asset / Content", "Provided By", "Notes"],
               [[name, asset_label(by), ""] for name, by in asset_items],
               [4000, 3000, 2360])
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 11 — POST-DELIVERY & SUPPORT
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "11.  POST-DELIVERY & SUPPORT")
    spacer(doc, 40, 40)

    sc_owner     = data.get("sourceCodeOwner", "")
    hosting_aft  = data.get("hostingAfter",    "")
    maintenance  = data.get("maintenance",     "")
    maint_fee    = data.get("maintenanceFee",  "")

    def support_fn(cell):
        sp = data.get("supportPeriod", "")
        lv_para(cell, "Free Support Period After Delivery",
                f"{sp} days (bug fixes only — no new features)" if sp else "")
        scp = cell.add_paragraph()
        add_paragraph_spacing(scp, before=0, after=56)
        add_run(scp, "Source Code Handover:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        add_run(scp, f"{cb(sc_owner == 'included')}  Included      "
                     f"{cb(sc_owner == 'not-included')}  Not Included      "
                     f"{cb(sc_owner == 'extra')}  Available at Extra Cost",
                size_pt=11, color=DARK_GRAY)
        hp = cell.add_paragraph()
        add_paragraph_spacing(hp, before=0, after=56)
        add_run(hp, "Hosting Management After Delivery:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        add_run(hp, f"{cb(hosting_aft == 'codesino')}  Managed by Codesino      "
                    f"{cb(hosting_aft == 'client')}  Transferred to Client      "
                    f"{cb(hosting_aft == 'tbd')}  TBD",
                size_pt=11, color=DARK_GRAY)
        mp = cell.add_paragraph()
        no_space(mp)
        add_run(mp, "Ongoing Maintenance Retainer:  ", bold=True, size_pt=11, color=BRAND_BLUE)
        if maintenance == "agreed":
            add_run(mp, f"☑  Agreed — {currency_sym}{maint_fee} / month      ☐  Not Agreed",
                    size_pt=11, color=DARK_GRAY)
        else:
            add_run(mp, "☐  Agreed      ☑  Not Agreed", size_pt=11, color=DARK_GRAY)

    section_box(doc, support_fn, HEX_WHITE)
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 12 — INTELLECTUAL PROPERTY
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "12.  INTELLECTUAL PROPERTY")
    for item in [
        "Upon receipt of full and final payment, all intellectual property rights for the final deliverables transfer entirely to the client.",
        "Codesino retains the right to display the completed project in its portfolio and marketing materials unless the client requests otherwise in writing prior to project completion.",
        "Any third-party libraries, plugins, frameworks, or open-source software used remain subject to their own respective licences.",
        "Codesino retains full ownership of all custom code, internal tools, and reusable components used in the project until final payment is received and confirmed.",
        "The client warrants that all content, images, logos, and materials provided to Codesino are legally owned or licenced by the client and do not infringe any third-party intellectual property rights.",
    ]:
        add_numbered(doc, item)
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 13 — CONFIDENTIALITY
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "13.  CONFIDENTIALITY")
    body_text(doc,
              "Both parties agree to keep strictly confidential any sensitive business information, "
              "trade secrets, proprietary data, client lists, pricing, or technical specifications "
              "shared during the course of this project. This obligation remains in full effect for "
              "two (2) years after the completion or termination of this agreement.")
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 14 — TERMINATION
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "14.  TERMINATION")
    for item in [
        "Either party may terminate this agreement with 7 (seven) calendar days' written notice.",
        "If the client terminates the agreement after development has commenced, all payments made are non-refundable.",
        "If Codesino terminates the agreement without cause, a pro-rated refund will be issued for work not yet started.",
        "Termination must be communicated via email or documented written message to be legally valid.",
        "Upon termination, each party shall immediately return or destroy any confidential materials belonging to the other party.",
    ]:
        add_numbered(doc, item)
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 15 — DISPUTE RESOLUTION
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "15.  DISPUTE RESOLUTION")
    body_text(doc,
              "In the event of a dispute, both parties agree to first attempt resolution through "
              "good-faith negotiation within 14 (fourteen) calendar days of the dispute being raised "
              "in writing. If resolution cannot be reached, both parties agree to submit the matter "
              "to mediation before resorting to formal legal proceedings. This agreement shall be "
              "governed by and construed in accordance with the laws of the Federal Republic of Nigeria.")
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 16 — ADDITIONAL NOTES
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "16.  ADDITIONAL NOTES & SPECIAL CONDITIONS")
    notes = data.get("additionalNotes", "")
    if notes:
        body_text(doc, notes)
    else:
        italic_note(doc, "No additional notes for this agreement.")
    divider(doc)

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION 17 — SIGNATURES
    # ═════════════════════════════════════════════════════════════════════════
    heading1(doc, "17.  AGREEMENT & SIGNATURES")
    body_text(doc,
              "By signing below, both parties confirm that they have read, understood, and agreed "
              "to all terms and conditions outlined in this Client Project Agreement. This document "
              "constitutes a legally binding agreement between Codesino Software Development "
              "Services and the client named herein.")
    spacer(doc, 80, 60)

    # ── Signature table: [Provider block] [gap] [Client block] ──────────────
    sig_table = doc.add_table(rows=2, cols=3)
    set_table_width(sig_table,    9360)
    set_column_widths(sig_table, [4400, 560, 4400])
    set_table_no_borders(sig_table)

    # Header row — coloured title bands
    for col_idx, label, align in [
        (0, "SERVICE PROVIDER",  WD_ALIGN_PARAGRAPH.CENTER),
        (2, "CLIENT",            WD_ALIGN_PARAGRAPH.CENTER),
    ]:
        cell = sig_table.cell(0, col_idx)
        set_cell_shading(cell,  HEX_BRAND_BLUE)
        set_cell_margins(cell,  top=100, bottom=100, left=150, right=150)
        set_cell_borders(cell,  top=None, bottom=None, left=None, right=None)
        clear_cell_paras(cell)
        p = cell.add_paragraph()
        no_space(p)
        p.alignment = align
        add_run(p, label, bold=True, size_pt=11, color=WHITE_COLOR)

    # Spacer column header
    gap_hdr = sig_table.cell(0, 1)
    set_cell_borders(gap_hdr, top=None, bottom=None, left=None, right=None)
    clear_cell_paras(gap_hdr)
    gap_hdr.add_paragraph()

    # Body row — signature content cells
    thin = {"val": "single", "sz": 2, "color": HEX_RULE}

    # Provider signature cell
    prov_cell = sig_table.cell(1, 0)
    set_cell_borders(prov_cell, top=None, bottom=thin, left=thin, right=thin)
    set_cell_margins(prov_cell, top=120, bottom=140, left=150, right=150)
    clear_cell_paras(prov_cell)

    pp1 = prov_cell.add_paragraph()
    add_paragraph_spacing(pp1, before=0, after=50)
    add_run(pp1, "Codesino Software Development Services",
            size_pt=10, color=DARK_GRAY)

    pp2 = prov_cell.add_paragraph()
    add_paragraph_spacing(pp2, before=0, after=80)
    add_run(pp2, "Authorised Signatory", size_pt=10, color=MID_GRAY)

    sig_line(prov_cell, "Digital Signature:")

    # ── Stamp image: placed below the sig line inside the provider cell ──────
    if os.path.exists(STAMP_PATH):
        stamp_p = prov_cell.add_paragraph()
        add_paragraph_spacing(stamp_p, before=20, after=20)
        stamp_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        stamp_p.add_run().add_picture(STAMP_PATH, width=Inches(1.3))

    pds  = data.get("providerSignDate", "")
    pp4  = prov_cell.add_paragraph()
    no_space(pp4)
    add_run(pp4, f"Date: {format_date(pds) if pds else '___ / ___ / ______'}",
            italic=True, size_pt=10, color=hex_to_rgb("999999"))

    # Spacer body cell
    gap_body = sig_table.cell(1, 1)
    set_cell_borders(gap_body, top=None, bottom=None, left=None, right=None)
    clear_cell_paras(gap_body)
    gap_body.add_paragraph()

    # Client signature cell
    cli_cell = sig_table.cell(1, 2)
    set_cell_borders(cli_cell, top=None, bottom=thin, left=thin, right=thin)
    set_cell_margins(cli_cell, top=120, bottom=140, left=150, right=150)
    clear_cell_paras(cli_cell)

    cc = data.get("clientCompany", "")
    cn = data.get("clientName",    "")

    cp1 = cli_cell.add_paragraph()
    add_paragraph_spacing(cp1, before=0, after=50)
    add_run(cp1, cc if cc else "____________________________",
            italic=not bool(cc), size_pt=10,
            color=DARK_GRAY if cc else hex_to_rgb("999999"))

    cp2 = cli_cell.add_paragraph()
    add_paragraph_spacing(cp2, before=0, after=80)
    add_run(cp2, f"Name: {cn}" if cn else "Name: ______________________",
            italic=not bool(cn), size_pt=10,
            color=DARK_GRAY if cn else hex_to_rgb("999999"))

    sig_line(cli_cell, "Signature:")

    # Blank space where client stamps/signs (no stamp image for client)
    blank_p = cli_cell.add_paragraph()
    add_paragraph_spacing(blank_p, before=20, after=20)
    no_space(blank_p)

    cds  = data.get("clientSignDate", "")
    cp4  = cli_cell.add_paragraph()
    no_space(cp4)
    add_run(cp4, f"Date: {format_date(cds) if cds else '___ / ___ / ______'}",
            italic=True, size_pt=10, color=hex_to_rgb("999999"))

    spacer(doc, 100, 60)

    # ── Document footer text ─────────────────────────────────────────────────
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(fp)
    add_run(fp,
            "This agreement was prepared by Codesino Software Development Services  |  "
            "www.codesinodev.com  |  contact@codesinodev.com",
            italic=True, size_pt=9, color=hex_to_rgb("9E9E9E"))

    # ── Page-level running footer ────────────────────────────────────────────
    section = doc.sections[0]
    footer  = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_para,
            "Codesino Software Development Services  |  Client Project Agreement  |  CONFIDENTIAL",
            size_pt=8, color=hex_to_rgb("999999"))

    # ── Serialize to bytes ───────────────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════
# GENERATOR 2 — EXECUTIVE / SENIOR ROLES
# ═══════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────

def generate_executive_agreement_buffer(data: dict) -> bytes:
    """
    Endpoint — Employment Agreement for Executive / Senior Company Roles.
    Positions: CTO, Co-Founder, Manager, etc.
    """
    doc            = _setup_document()
    ref            = generate_ref("EXC")
    date_str       = format_datetime()
    currency_sym   = "$" if data.get("currency") == "usd" else "₦"
    employee_name  = data.get("employeeName", "the incoming team member")
    position       = data.get("position", "the stated position")

    _build_header(doc, "Executive & Senior Staff Employment Agreement", ref, date_str)

    # ─── CONGRATULATION OPENING ───────────────────────────────────────────
    heading1(doc, "1.  LETTER OF APPOINTMENT")
    spacer(doc, 40, 40)

    def opening_fn(cell):
        p1 = cell.add_paragraph()
        add_paragraph_spacing(p1, before=0, after=60)
        add_run(p1, f"Dear {employee_name},", bold=True, size_pt=12, color=BRAND_BLUE)
        p2 = cell.add_paragraph()
        add_paragraph_spacing(p2, before=0, after=60, line=276)
        add_run(p2,
                f"On behalf of everyone at Codesino Software Development Services, "
                f"we are truly delighted to extend this formal offer of employment to you. "
                f"Your skills, experience, and the impression you made throughout our "
                f"interview and evaluation process made this decision straightforward — "
                f"we believe you are the right person for this role and for this company.",
                size_pt=11, color=DARK_GRAY)
        p3 = cell.add_paragraph()
        add_paragraph_spacing(p3, before=0, after=60, line=276)
        add_run(p3,
                f"You are hereby appointed to the position of "
                f"\"{v(position, 'the stated senior role')}\" "
                f"at Codesino Software Development Services, effective from the date of hire "
                f"as specified in Section 2 of this agreement. This is a senior leadership "
                f"position that carries significant responsibility, and we are confident in "
                f"your ability to rise to that challenge and make a lasting impact.",
                size_pt=11, color=DARK_GRAY)
        p4 = cell.add_paragraph()
        add_paragraph_spacing(p4, before=0, after=0, line=276)
        add_run(p4,
                "Welcome to the Codesino family. We look forward to building something "
                "extraordinary together.",
                size_pt=11, color=DARK_GRAY)

    section_box(doc, opening_fn, HEX_LIGHT_BLUE, accent_color=HEX_MID_BLUE)
    divider(doc)

    # ─── EMPLOYMENT DETAILS ────────────────────────────────────────────────
    _build_employment_details(doc, data, "Position / Role",
                              EXECUTIVE_POSITIONS, allow_multi_position=False)

    # ─── COMPENSATION ─────────────────────────────────────────────────────
    _build_compensation(doc, data, currency_sym)

    # ─── RESPONSIBILITIES ─────────────────────────────────────────────────
    heading1(doc, "4.  KEY RESPONSIBILITIES")
    italic_note(doc,
                "The responsibilities listed below are associated with this senior position. "
                "Additional duties may be assigned at the discretion of management.")
    spacer(doc, 40, 40)

    default_exec_duties = [
        "Provide strategic leadership and direction for your assigned department or function.",
        "Collaborate with the founding team and senior management to define and execute company goals.",
        "Manage, mentor, and hold accountable the team members under your supervision.",
        "Make high-level decisions that positively impact the company's operations, culture, and output.",
        "Represent Codesino with professionalism at all times — both internally and externally.",
        "Participate actively in company planning, performance reviews, and strategic discussions.",
        "Ensure all activities within your purview meet Codesino's quality, ethical, and compliance standards.",
        "Report progress, blockers, and performance metrics to executive management on an agreed schedule.",
    ]
    extra_duties = [d for d in (data.get("additionalDuties") or []) if d]
    all_duties   = default_exec_duties + extra_duties
    for duty in all_duties:
        add_bullet(doc, duty)

    if data.get("specificResponsibilities"):
        spacer(doc, 80, 40)
        heading2(doc, "Role-Specific Responsibilities")
        body_text(doc, data.get("specificResponsibilities", ""))

    divider(doc)

    # ─── COMPANY NOTICE & CONDUCT ─────────────────────────────────────────
    heading1(doc, "5.  COMPANY NOTICE, CONDUCT & OBLIGATIONS")
    spacer(doc, 40, 40)

    def notice_fn(cell):
        p0 = cell.add_paragraph()
        add_paragraph_spacing(p0, before=0, after=60)
        add_run(p0, "IMPORTANT NOTICE TO THE EMPLOYEE", bold=True, size_pt=12, color=BRAND_BLUE)

        p1 = cell.add_paragraph()
        add_paragraph_spacing(p1, before=0, after=50, line=276)
        add_run(p1,
                "As a senior member of the Codesino Software Development Services team, "
                "you occupy a position of considerable trust, access, and influence. "
                "By accepting this employment offer, you expressly acknowledge and agree to "
                "abide by the following obligations without exception.",
                size_pt=11, color=DARK_GRAY)

    section_box(doc, notice_fn, HEX_PALE_BLUE, accent_color=HEX_BRAND_BLUE)
    spacer(doc, 60, 40)

    heading2(doc, "5.1  Acceptance of Company Policies")
    body_text(doc,
              "By accepting this appointment, you confirm that you have read, understood, and "
              "agree to comply fully with Codesino Software Development Services' Privacy Policy "
              "and Terms of Service, both of which are integral components of your employment "
              "contract. These documents govern how the company handles data, client interactions, "
              "business operations, and internal conduct. Failure to abide by these policies "
              "constitutes a breach of this agreement.")

    heading2(doc, "5.2  Strict Confidentiality of Company Information")
    for item in [
        "You shall treat all internal company information — including but not limited to "
        "business strategies, financial data, product roadmaps, pricing structures, internal "
        "communications, team compositions, operational processes, proprietary tools, and "
        "source code — as strictly confidential.",
        "This obligation applies both during and after the term of your employment with "
        "Codesino. You shall not, at any time, disclose, share, publish, or allow access to "
        "any such information to any unauthorised party.",
        "Any deliberate or negligent exposure of company-confidential information will be "
        "treated as gross misconduct and may result in immediate termination of employment, "
        "civil legal action, and/or criminal proceedings, depending on the nature and severity "
        "of the breach.",
        "You must immediately report any suspected or actual breach of company information "
        "security to the appropriate authority within the company, regardless of who the "
        "perpetrator may be.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.3  Strict Confidentiality of Client Information")
    for item in [
        "As a senior staff member, you will likely have access to client details, project "
        "briefs, payment information, personal data, business strategies, and communications. "
        "All such client information is to be treated with the highest level of confidentiality.",
        "You shall never share, discuss, reference, or use client information outside of the "
        "direct scope of work for which that information was provided. This applies to verbal, "
        "written, digital, and indirect disclosures.",
        "Client confidentiality extends to all platforms, including but not limited to social "
        "media, personal blogs, public forums, casual conversations, networking events, "
        "and third-party communication platforms.",
        "Any breach of client confidentiality — whether intentional or otherwise — will result "
        "in immediate termination of employment, financial penalties, and potential civil or "
        "criminal legal action against the employee. The company will not hesitate to pursue "
        "full legal remedies to protect its clients.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.4  Professional Conduct & Ethical Standards")
    for item in [
        "As a leader within this organisation, you are expected to model exemplary professional "
        "conduct at all times — in interactions with colleagues, clients, vendors, and the "
        "public — both in professional settings and on personal platforms.",
        "You shall not engage in, condone, or facilitate any activity that is dishonest, "
        "discriminatory, exploitative, or harmful to the company, its employees, or its clients.",
        "You shall not enter into any external business arrangements, partnerships, or "
        "employment that creates a direct or indirect conflict of interest with Codesino "
        "Software Development Services without prior written consent from the company.",
        "Any form of insider trading, corporate espionage, solicitation of the company's "
        "clients or team members for competing ventures, or misappropriation of company "
        "resources is strictly prohibited and will be prosecuted to the fullest extent of the law.",
        "You are responsible for maintaining a healthy, respectful, and productive work "
        "environment. Any form of workplace harassment, bullying, or discrimination will not "
        "be tolerated and may constitute grounds for immediate dismissal.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.5  Non-Solicitation & Non-Compete")
    for item in [
        "During your employment and for a period of 12 (twelve) months following the "
        "termination of this agreement — for any reason — you shall not directly or "
        "indirectly solicit, approach, or recruit any of the company's employees, contractors, "
        "or clients for the benefit of a competing business.",
        "You shall not use any proprietary knowledge, trade secrets, client relationships, "
        "or internal systems knowledge gained at Codesino to establish or assist a directly "
        "competing business within the same industry and market.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.6  Additional Company Rules")
    additional_rules = [
        "All work produced during the course of your employment — including ideas, designs, "
        "strategies, and creative output — belongs exclusively to Codesino Software "
        "Development Services.",
        "You are required to maintain professional development and remain current with "
        "developments in your field of expertise relevant to your role.",
        "All company devices, credentials, systems, and digital assets provided to you "
        "remain the property of the company and must be used exclusively for company-related "
        "purposes unless explicitly stated otherwise.",
        "Any reputational damage caused by an employee's public actions or statements — "
        "whether online or offline — will be taken seriously and may constitute grounds for "
        "disciplinary action, up to and including termination.",
    ]
    for rule in additional_rules:
        add_numbered(doc, rule)

    if data.get("additionalNotice"):
        spacer(doc, 60, 40)
        heading2(doc, "5.7  Specific Instructions from Management")
        body_text(doc, data.get("additionalNotice", ""))

    divider(doc)

    # ─── ACCEPTANCE STATEMENT ─────────────────────────────────────────────
    heading1(doc, "6.  ACCEPTANCE OF EMPLOYMENT TERMS")
    spacer(doc, 40, 40)

    def acceptance_fn(cell):
        p = cell.add_paragraph()
        add_paragraph_spacing(p, before=0, after=60, line=276)
        add_run(p,
                "By accepting this employment offer and appending your signature to this "
                "document, you confirm that:",
                bold=True, size_pt=11, color=BRAND_BLUE)
        for item in [
            "You have read and understood the full contents of this Employment Agreement.",
            "You agree to abide by all company policies, privacy obligations, confidentiality "
            "requirements, and conduct standards outlined herein.",
            "You understand that any breach of these obligations may result in immediate "
            "termination of employment and legal action.",
            "You accept the compensation, responsibilities, and conditions of employment "
            "as described in this agreement.",
            "You agree that all work produced during your employment is the exclusive "
            "intellectual property of Codesino Software Development Services.",
        ]:
            add_bullet(cell, item)
        pn = cell.add_paragraph()
        add_paragraph_spacing(pn, before=60, after=0)
        add_run(pn,
                "This agreement takes effect from the date of hire as specified in Section 2 "
                "and remains in force until lawfully terminated by either party.",
                italic=True, size_pt=10, color=MUTED_BLUE)

    section_box(doc, acceptance_fn, HEX_LIGHT_BLUE, accent_color=HEX_MID_BLUE)
    divider(doc)

    # ─── IP + STANDARD CLAUSES ────────────────────────────────────────────
    _build_standard_clauses(doc)

    # ─── ADDITIONAL NOTES ─────────────────────────────────────────────────
    heading1(doc, "ADDITIONAL NOTES & SPECIAL CONDITIONS")
    notes = data.get("additionalNotes", "")
    if notes:
        body_text(doc, notes)
    else:
        italic_note(doc, "No additional notes for this agreement.")
    divider(doc)

    # ─── SIGNATURES ───────────────────────────────────────────────────────
    _build_signatures(doc, data)
    _build_footer(doc, "Executive Employment Agreement")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════
# GENERATOR 3 — SUPPORT / OPERATIONS ROLES
# ═══════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────

def generate_support_agreement_buffer(data: dict) -> bytes:
    """
    Endpoint — Employment Agreement for Support & Operations Roles.
    Positions: Customer Support, Social Media, etc.
    Employees may hold more than one position.
    """
    doc            = _setup_document()
    ref            = generate_ref("SUP")
    date_str       = format_datetime()
    currency_sym   = "$" if data.get("currency") == "usd" else "₦"
    employee_name  = data.get("employeeName", "the incoming team member")

    selected_pos   = data.get("positions", [])
    if isinstance(selected_pos, str):
        selected_pos = [selected_pos]
    pos_display    = ", ".join(selected_pos) if selected_pos else "the stated role(s)"

    _build_header(doc, "Support & Operations Staff Employment Agreement", ref, date_str)

    # ─── CONGRATULATION OPENING ───────────────────────────────────────────
    heading1(doc, "1.  LETTER OF APPOINTMENT")
    spacer(doc, 40, 40)

    def opening_fn(cell):
        p1 = cell.add_paragraph()
        add_paragraph_spacing(p1, before=0, after=60)
        add_run(p1, f"Dear {employee_name},", bold=True, size_pt=12, color=BRAND_BLUE)
        p2 = cell.add_paragraph()
        add_paragraph_spacing(p2, before=0, after=60, line=276)
        add_run(p2,
                "We are excited to welcome you to the Codesino Software Development Services "
                "team! After careful consideration, we are pleased to confirm your appointment "
                "to the Codesino support and operations team. Your dedication, communication "
                "skills, and readiness to grow with the team impressed us, and we are looking "
                "forward to the contributions you will make.",
                size_pt=11, color=DARK_GRAY)
        p3 = cell.add_paragraph()
        add_paragraph_spacing(p3, before=0, after=60, line=276)
        add_run(p3,
                f"You are hereby appointed to the role(s) of \"{pos_display}\" within "
                f"the Support & Operations division at Codesino Software Development Services, "
                f"effective from the date of hire specified in Section 2 of this agreement. "
                f"Please review this agreement carefully before signing.",
                size_pt=11, color=DARK_GRAY)
        p4 = cell.add_paragraph()
        add_paragraph_spacing(p4, before=0, after=0, line=276)
        add_run(p4,
                "Welcome aboard — we are glad to have you on the team!",
                size_pt=11, color=DARK_GRAY)

    section_box(doc, opening_fn, HEX_LIGHT_BLUE, accent_color=HEX_MID_BLUE)
    divider(doc)

    # ─── EMPLOYMENT DETAILS (multi-position) ──────────────────────────────
    _build_employment_details(doc, data, "Position(s) / Role(s)",
                              SUPPORT_POSITIONS, allow_multi_position=True)

    # ─── COMPENSATION ─────────────────────────────────────────────────────
    _build_compensation(doc, data, currency_sym)

    # ─── RESPONSIBILITIES ─────────────────────────────────────────────────
    heading1(doc, "4.  KEY RESPONSIBILITIES")
    italic_note(doc,
                "The responsibilities below apply to all support and operations staff. "
                "Role-specific duties are outlined separately beneath.")
    spacer(doc, 40, 40)

    heading2(doc, "4.1  General Responsibilities")
    general_support_duties = [
        "Represent Codesino Software Development Services with professionalism, "
        "warmth, and competence in every interaction — both internally and externally.",
        "Attend all scheduled team check-ins, briefings, and communication sessions "
        "punctually and come prepared.",
        "Complete all assigned tasks within agreed deadlines and maintain a consistent "
        "quality of output at all times.",
        "Proactively communicate blockers, delays, or concerns to the relevant team lead "
        "or supervisor in a timely manner.",
        "Maintain all company tools, accounts, and platforms assigned to your role "
        "with care and only use them for company-related purposes.",
        "Support team members across departments as needed and contribute to a positive, "
        "collaborative workplace culture.",
    ]
    for duty in general_support_duties:
        add_bullet(doc, duty)

    spacer(doc, 80, 40)
    heading2(doc, "4.2  Customer Support — Role-Specific Responsibilities")
    italic_note(doc,
                "The following apply specifically to employees serving in a Customer Support capacity.")
    cs_duties = [
        "Respond to all customer enquiries, complaints, and requests with empathy, patience, "
        "and professionalism. Customers must always feel heard, respected, and valued — "
        "regardless of the situation.",
        "Maintain outstanding communication standards at every touchpoint. Your language, "
        "tone, and response time are a direct reflection of the Codesino brand.",
        "Accurately document all customer issues, feedback, and complaints in the designated "
        "tracking or ticketing system in real time. No unresolved issue should go unlogged.",
        "Escalate complex or technical issues to the appropriate internal reporting team "
        "promptly, with a clear summary of the problem, the customer's details, and any "
        "steps already taken. Clear handover notes are non-negotiable.",
        "Follow up diligently with customers after an issue has been escalated or resolved "
        "to ensure their satisfaction and close the loop professionally.",
        "Identify recurring customer pain points and report them to your supervisor as "
        "part of ongoing quality assurance efforts.",
        "Never make promises to customers about features, timelines, or outcomes that have "
        "not been confirmed by the technical or management team.",
        "Treat every customer interaction as an opportunity to build long-term trust in "
        "the Codesino brand — because satisfied customers become loyal advocates.",
    ]
    for duty in cs_duties:
        add_numbered(doc, duty)

    spacer(doc, 80, 40)
    heading2(doc, "4.3  Social Media & Content — Role-Specific Responsibilities")
    italic_note(doc,
                "The following apply specifically to employees in Social Media, Content, or Community roles.")
    sm_duties = [
        "Manage all assigned Codesino social media platforms with consistency, creativity, "
        "and a clear understanding of the company's brand voice, tone, and target audience.",
        "Create, schedule, and publish content that is engaging, accurate, and aligned "
        "with the company's marketing strategy and brand guidelines.",
        "Monitor all channels for comments, messages, and mentions. Respond promptly and "
        "professionally, and escalate any sensitive matters to the relevant supervisor.",
        "Track and report on content performance metrics regularly — including engagement "
        "rates, reach, follower growth, and campaign results — and make data-informed "
        "recommendations for improvement.",
        "Collaborate with the design and development teams to ensure all social content "
        "is visually on-brand and technically accurate.",
        "Never publish content about clients, projects, or internal company matters "
        "without explicit written approval from management.",
        "Stay current with social media trends, platform updates, and best practices, "
        "and proactively bring relevant ideas to the team.",
    ]
    for duty in sm_duties:
        add_numbered(doc, duty)

    if data.get("specificResponsibilities"):
        spacer(doc, 80, 40)
        heading2(doc, "4.4  Additional Role-Specific Instructions")
        body_text(doc, data.get("specificResponsibilities", ""))

    if data.get("additionalDuties"):
        spacer(doc, 60, 40)
        heading2(doc, "4.5  Other Assigned Duties")
        for duty in (data.get("additionalDuties") or []):
            if duty:
                add_bullet(doc, duty)

    divider(doc)

    # ─── COMPANY NOTICE & CONDUCT ─────────────────────────────────────────
    heading1(doc, "5.  COMPANY NOTICE, CONDUCT & OBLIGATIONS")
    spacer(doc, 40, 40)

    def notice_fn(cell):
        p0 = cell.add_paragraph()
        add_paragraph_spacing(p0, before=0, after=60)
        add_run(p0, "IMPORTANT NOTICE TO THE EMPLOYEE", bold=True, size_pt=12, color=BRAND_BLUE)
        p1 = cell.add_paragraph()
        add_paragraph_spacing(p1, before=0, after=50, line=276)
        add_run(p1,
                "As a member of the Codesino support and operations team, you are the front "
                "line of our brand. Your interactions with customers, the public, and internal "
                "teams directly shape the reputation and integrity of this company. "
                "Please read and understand the following obligations thoroughly.",
                size_pt=11, color=DARK_GRAY)

    section_box(doc, notice_fn, HEX_PALE_BLUE, accent_color=HEX_BRAND_BLUE)
    spacer(doc, 60, 40)

    heading2(doc, "5.1  Acceptance of Company Policies")
    body_text(doc,
              "By accepting this appointment, you confirm that you have read and agree to "
              "comply fully with Codesino Software Development Services' Privacy Policy and "
              "Terms of Service. These policies govern how you handle data, client "
              "interactions, and your conduct as a representative of the company.")

    heading2(doc, "5.2  Confidentiality of Company Information")
    for item in [
        "All company-internal information — including business processes, internal tools, "
        "team structures, financial data, client lists, campaign strategies, and performance "
        "metrics — is strictly confidential and must not be shared with anyone outside the company.",
        "This obligation applies during your employment and continues for a period of "
        "2 (two) years following the termination of this agreement.",
        "Sharing confidential company information through any medium — including social media, "
        "messaging applications, email, or verbal communication — without explicit authorisation "
        "is a serious breach of this agreement and will result in immediate termination "
        "and potential legal action.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.3  Confidentiality of Client Information")
    for item in [
        "In the course of your role, you may have access to client details, orders, project "
        "information, personal data, or communications. All such information is strictly "
        "confidential and must be handled with the utmost discretion.",
        "Client information must never be discussed, shared, posted, or referenced outside "
        "of your direct work responsibilities — not on social media, not with friends or "
        "family, and not in any other public or private setting.",
        "Any unauthorised disclosure of client information — whether intentional or "
        "accidental — will result in immediate termination of employment and may expose "
        "the employee to civil or criminal liability. The company will actively pursue "
        "all available legal remedies to protect its clients.",
        "If you suspect that client data has been compromised or exposed, you must "
        "report this to your supervisor immediately.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.4  Customer Interaction Standards")
    for item in [
        "Every customer — regardless of the nature of their concern — deserves to be "
        "treated with patience, respect, and genuine care. There are no exceptions to this standard.",
        "Responses to customers must always be professional, clear, accurate, and timely. "
        "Poor communication is not acceptable at any level.",
        "Never engage in heated exchanges, dismissive language, or unprofessional behaviour "
        "with customers under any circumstances. If a situation becomes difficult, escalate "
        "it to a supervisor rather than engaging further.",
        "You are not authorised to offer refunds, discounts, or make commitments on behalf "
        "of the company without prior written approval from a manager.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.5  Online & Social Media Conduct")
    for item in [
        "As a staff member, your public online presence is subject to conduct standards "
        "that protect the company's reputation. You must not post, share, or publish "
        "any content that disparages Codesino, its clients, or its team members.",
        "You must not represent yourself as an official spokesperson for Codesino on "
        "personal platforms unless explicitly authorised to do so in writing.",
        "Any content related to company projects, clients, or internal operations "
        "requires management approval before publication — regardless of the platform.",
    ]:
        add_numbered(doc, item)

    if data.get("additionalNotice"):
        spacer(doc, 60, 40)
        heading2(doc, "5.6  Specific Instructions from Management")
        body_text(doc, data.get("additionalNotice", ""))

    divider(doc)

    # ─── ACCEPTANCE STATEMENT ─────────────────────────────────────────────
    heading1(doc, "6.  ACCEPTANCE OF EMPLOYMENT TERMS")
    spacer(doc, 40, 40)

    def acceptance_fn(cell):
        p = cell.add_paragraph()
        add_paragraph_spacing(p, before=0, after=60, line=276)
        add_run(p,
                "By accepting this employment and signing below, I confirm that:",
                bold=True, size_pt=11, color=BRAND_BLUE)
        for item in [
            "I have read and fully understood this Employment Agreement.",
            "I agree to abide by all conduct standards, confidentiality obligations, "
            "and company policies outlined in this document.",
            "I understand that breaches of confidentiality or misconduct may result in "
            "immediate termination of employment and legal consequences.",
            "I accept the compensation structure, roles, and conditions of employment "
            "as described in this agreement.",
        ]:
            add_bullet(cell, item)
        pn = cell.add_paragraph()
        add_paragraph_spacing(pn, before=60, after=0)
        add_run(pn,
                "This agreement is effective from the date of hire specified in Section 2.",
                italic=True, size_pt=10, color=MUTED_BLUE)

    section_box(doc, acceptance_fn, HEX_LIGHT_BLUE, accent_color=HEX_MID_BLUE)
    divider(doc)

    _build_standard_clauses(doc)

    heading1(doc, "ADDITIONAL NOTES & SPECIAL CONDITIONS")
    notes = data.get("additionalNotes", "")
    if notes:
        body_text(doc, notes)
    else:
        italic_note(doc, "No additional notes for this agreement.")
    divider(doc)

    _build_signatures(doc, data)
    _build_footer(doc, "Support & Operations Employment Agreement")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════
# GENERATOR 4 — TECHNICAL / DEVELOPMENT TEAM ROLES
# ═══════════════════════════════════════════════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────

def generate_technical_agreement_buffer(data: dict) -> bytes:
    """
    Endpoint — Employment Agreement for Technical / Dev Team Roles.
    Positions: Frontend, Backend, DevOps, UI/UX, etc.
    """
    doc            = _setup_document()
    ref            = generate_ref("TEC")
    date_str       = format_datetime()
    currency_sym   = "$" if data.get("currency") == "usd" else "₦"
    employee_name  = data.get("employeeName", "the incoming team member")
    position       = data.get("position", "the stated technical role")

    _build_header(doc, "Technical Team Employment Agreement", ref, date_str)

    # ─── CONGRATULATION OPENING ───────────────────────────────────────────
    heading1(doc, "1.  LETTER OF APPOINTMENT")
    spacer(doc, 40, 40)

    def opening_fn(cell):
        p1 = cell.add_paragraph()
        add_paragraph_spacing(p1, before=0, after=60)
        add_run(p1, f"Dear {employee_name},", bold=True, size_pt=12, color=BRAND_BLUE)
        p2 = cell.add_paragraph()
        add_paragraph_spacing(p2, before=0, after=60, line=276)
        add_run(p2,
                "Congratulations — and a very warm welcome to the Codesino Software Development "
                "Services engineering team! We are genuinely excited about the talent, "
                "perspective, and technical ability you bring, and we are confident that your "
                "work here will have a real and lasting impact on what we build and how we "
                "build it.",
                size_pt=11, color=DARK_GRAY)
        p3 = cell.add_paragraph()
        add_paragraph_spacing(p3, before=0, after=60, line=276)
        add_run(p3,
                f"You are hereby appointed to the position of "
                f"\"{v(position, 'the stated technical role')}\" "
                f"within the Codesino Technical Team, effective from the date of hire "
                f"specified in Section 2 of this agreement. We chose you because of your "
                f"demonstrated skill, your problem-solving mindset, and the quality of your "
                f"technical thinking — and we look forward to seeing you thrive.",
                size_pt=11, color=DARK_GRAY)
        p4 = cell.add_paragraph()
        add_paragraph_spacing(p4, before=0, after=0, line=276)
        add_run(p4,
                "Welcome to the build team. Let's create something great together.",
                size_pt=11, color=DARK_GRAY)

    section_box(doc, opening_fn, HEX_LIGHT_BLUE, accent_color=HEX_MID_BLUE)
    divider(doc)

    # ─── EMPLOYMENT DETAILS ────────────────────────────────────────────────
    _build_employment_details(doc, data, "Technical Position / Role",
                              TECHNICAL_POSITIONS, allow_multi_position=False)

    # Extra technical fields
    heading2(doc, "2.2  Technical Stack / Tools")

    def stack_fn(cell):
        lv_para(cell, "Primary Language(s) / Framework(s)", data.get("primaryStack",   ""))
        lv_para(cell, "Secondary Tools / Technologies",     data.get("secondaryStack", ""))
        lv_para(cell, "Assigned Projects",                  data.get("assignedProjects",""))
        lv_para(cell, "Repository / Codebase Access",       data.get("repoAccess",     ""))
        lv_para(cell, "Preferred IDE / Tools",              data.get("preferredTools",  ""))

    section_box(doc, stack_fn, HEX_WHITE)
    divider(doc)

    # ─── COMPENSATION ─────────────────────────────────────────────────────
    _build_compensation(doc, data, currency_sym)

    # ─── RESPONSIBILITIES ─────────────────────────────────────────────────
    heading1(doc, "4.  KEY RESPONSIBILITIES & TECHNICAL STANDARDS")
    italic_note(doc,
                "All technical staff are expected to meet the following standards. "
                "Role-specific responsibilities are detailed in the subsections below.")
    spacer(doc, 40, 40)

    heading2(doc, "4.1  General Engineering Responsibilities")
    general_tech_duties = [
        "Write clean, well-structured, and thoroughly documented code that meets the "
        "company's code quality and review standards at all times.",
        "Participate in all code reviews, pull request approvals, and engineering "
        "discussions as part of the team's standard development workflow.",
        "Collaborate proactively with cross-functional teams — including designers, "
        "project managers, and other developers — to deliver cohesive, high-quality solutions.",
        "Report development progress, technical blockers, and timelines accurately "
        "and in a timely manner to the relevant project lead or manager.",
        "Adhere to all version control protocols and maintain clean, organised commits "
        "with clear and descriptive messages.",
        "Continuously improve your technical skills and stay current with industry tools, "
        "languages, frameworks, and best practices relevant to your role.",
        "Participate in retrospectives, post-mortems, and technical planning sessions "
        "as required by the team.",
    ]
    for duty in general_tech_duties:
        add_bullet(doc, duty)

    spacer(doc, 80, 40)
    heading2(doc, "4.2  Frontend Developer — Role-Specific Standards")
    italic_note(doc,
                "The following apply to all employees in Frontend Development or UI/UX Design roles.")
    fe_duties = [
        "You are responsible for building modern, polished, and highly intuitive user "
        "interfaces that represent the Codesino brand with excellence. Every pixel matters — "
        "visual quality and attention to detail are non-negotiable.",
        "All UI components must be fully responsive and thoroughly tested across all "
        "standard screen sizes and device types, from mobile to ultra-wide desktop. "
        "Mobile-first design is the default approach.",
        "Animations, transitions, and micro-interactions must be smooth, performant, "
        "and purposeful — enhancing the user experience without compromising performance "
        "or accessibility.",
        "Work closely with UX designers (or take ownership of UX when applicable) to "
        "ensure the end product is not just visually impressive but genuinely intuitive "
        "and enjoyable to use.",
        "All code must achieve high Lighthouse performance scores. Frontend performance "
        "is part of quality — slow interfaces are not acceptable deliverables.",
        "Ensure cross-browser compatibility across all major modern browsers: Chrome, "
        "Firefox, Safari, and Edge.",
        "Implement SEO-friendly markup, semantic HTML, and accessibility (a11y) best "
        "practices on all pages you build.",
        "Maintain a consistent design system and component library. Avoid one-off "
        "solutions where reusable components can be built instead.",
    ]
    for duty in fe_duties:
        add_numbered(doc, duty)

    spacer(doc, 80, 40)
    heading2(doc, "4.3  Backend Developer — Role-Specific Standards")
    italic_note(doc,
                "The following apply to all employees in Backend Development or API Engineering roles.")
    be_duties = [
        "Your primary mandate is to build robust, secure, and highly scalable backend "
        "systems that power the company's products and client solutions. Security is a "
        "first-class requirement — not an afterthought.",
        "All API endpoints must be thoroughly designed, documented, version-controlled, "
        "and tested before being handed off to the frontend team. Incomplete or "
        "undocumented APIs will not be accepted as deliverables.",
        "Implement enterprise-grade authentication and authorisation systems. Proper "
        "use of JWTs, OAuth, session management, and role-based access control is "
        "expected as a minimum standard.",
        "Apply rigorous input validation, parameterised queries, and rate limiting "
        "across all endpoints to prevent SQL injection, XSS, CSRF, and other common "
        "attack vectors. Building impregnable systems is a core expectation of this role.",
        "Design database schemas that are normalised, efficient, and future-proof. "
        "All database migrations must be scripted, versioned, and reversible.",
        "Write comprehensive unit tests and integration tests for all critical logic. "
        "Test coverage is not optional.",
        "Implement proper error handling, structured logging, and application monitoring "
        "to ensure systems can be diagnosed and recovered quickly when issues arise.",
        "Optimise all database queries, background jobs, and API response times. "
        "Backend performance directly impacts user experience and client satisfaction.",
    ]
    for duty in be_duties:
        add_numbered(doc, duty)

    spacer(doc, 80, 40)
    heading2(doc, "4.4  DevOps & Infrastructure — Role-Specific Standards")
    italic_note(doc,
                "The following apply to employees in DevOps, Cloud Infrastructure, or Platform Engineering roles.")
    devops_duties = [
        "Design, implement, and maintain reliable CI/CD pipelines that automate testing, "
        "building, and deployment processes across all environments.",
        "Manage and optimise all cloud infrastructure to ensure maximum availability, "
        "performance, and cost-efficiency. Infrastructure-as-Code (IaC) practices are "
        "expected at all times.",
        "Implement and enforce security hardening across all systems, networks, containers, "
        "and deployment environments.",
        "Monitor all production systems continuously and ensure incident response "
        "procedures are in place, documented, and tested.",
        "Maintain clear, up-to-date infrastructure documentation. All architecture "
        "decisions must be communicated to and approved by the appropriate stakeholders.",
        "Ensure backup, recovery, and disaster recovery protocols are in place and "
        "regularly tested for all critical systems.",
    ]
    for duty in devops_duties:
        add_numbered(doc, duty)

    if data.get("specificResponsibilities"):
        spacer(doc, 80, 40)
        heading2(doc, "4.5  Additional Role-Specific Instructions")
        body_text(doc, data.get("specificResponsibilities", ""))

    if data.get("additionalDuties"):
        spacer(doc, 60, 40)
        heading2(doc, "4.6  Other Assigned Duties")
        for duty in (data.get("additionalDuties") or []):
            if duty:
                add_bullet(doc, duty)

    divider(doc)

    # ─── COMPANY NOTICE & CONDUCT ─────────────────────────────────────────
    heading1(doc, "5.  COMPANY NOTICE, TECHNICAL OBLIGATIONS & CONDUCT")
    spacer(doc, 40, 40)

    def notice_fn(cell):
        p0 = cell.add_paragraph()
        add_paragraph_spacing(p0, before=0, after=60)
        add_run(p0, "IMPORTANT NOTICE TO THE EMPLOYEE", bold=True, size_pt=12, color=BRAND_BLUE)
        p1 = cell.add_paragraph()
        add_paragraph_spacing(p1, before=0, after=50, line=276)
        add_run(p1,
                "As a technical team member at Codesino Software Development Services, "
                "you will have access to proprietary source code, production systems, "
                "client data, and infrastructure credentials. The obligations outlined "
                "below are absolute — they are not guidelines but requirements.",
                size_pt=11, color=DARK_GRAY)

    section_box(doc, notice_fn, HEX_PALE_BLUE, accent_color=HEX_BRAND_BLUE)
    spacer(doc, 60, 40)

    heading2(doc, "5.1  Acceptance of Company Policies")
    body_text(doc,
              "By accepting this appointment, you confirm that you have read and agree to "
              "comply fully with Codesino Software Development Services' Privacy Policy and "
              "Terms of Service, which are binding on all technical team members and cover "
              "the handling of data, code, systems, and client information.")

    heading2(doc, "5.2  Source Code & Codebase Confidentiality")
    for item in [
        "All source code, codebases, technical architectures, database schemas, API "
        "designs, algorithms, and system configurations you work on or have access to "
        "are strictly proprietary and confidential assets of Codesino Software Development "
        "Services or its clients.",
        "You must not copy, export, share, publish, open-source, or otherwise distribute "
        "any portion of the company's or clients' codebases without explicit written "
        "authorisation from management. This applies even to code you personally authored "
        "during your employment.",
        "You must not use any proprietary code, architecture patterns, or technical "
        "solutions developed at Codesino for personal projects, freelance work, or "
        "third-party employers.",
        "If you leave the company for any reason, you must immediately revoke access "
        "to all company repositories, systems, and tools and confirm this in writing.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.3  Confidentiality of Company & Client Data")
    for item in [
        "In the course of your technical duties, you may have access to production databases, "
        "client records, user data, API keys, server credentials, and other sensitive system "
        "information. All such data is to be treated with the highest level of confidentiality.",
        "You must never export, store on personal devices, screenshot, or share any "
        "sensitive client or user data outside of the secure company environment, unless "
        "explicitly authorised and using company-approved methods.",
        "Any security vulnerability, data breach, or potential system compromise must be "
        "reported to the appropriate internal authority immediately — even if you were "
        "responsible for introducing the vulnerability. Prompt disclosure allows for "
        "mitigation; concealment is an unacceptable and terminable offence.",
        "Disclosing confidential company or client data to any external party — whether "
        "intentionally or through negligence — will result in immediate termination and "
        "active legal prosecution. The consequences are severe and the company will not "
        "negotiate in matters of data breach.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.4  Security Standards & Engineering Ethics")
    for item in [
        "You are expected to apply security best practices by default in everything you "
        "build. Security is not a feature to add later — it must be embedded from the "
        "very first line of code you write.",
        "You must not introduce backdoors, hidden functionality, hard-coded credentials, "
        "or any insecure patterns into company or client codebases under any circumstances.",
        "You must not use company systems, infrastructure, or client access to conduct "
        "personal research, run personal workloads, or carry out any activity not "
        "directly related to your assigned work.",
        "Attempting to access systems, data, or code repositories beyond your "
        "authorised access level — even out of curiosity — is a serious security "
        "violation and will be treated accordingly.",
    ]:
        add_numbered(doc, item)

    spacer(doc, 60, 40)
    heading2(doc, "5.5  Professional Conduct & Collaboration")
    for item in [
        "Engineering is a team sport at Codesino. You are expected to collaborate "
        "openly, communicate clearly, and support your colleagues constructively. "
        "Hoarding knowledge or building unnecessarily complex systems to create "
        "dependency is not acceptable.",
        "Code reviews are a learning and quality tool — give feedback respectfully "
        "and receive it graciously. All code review comments must be professional and "
        "focused on the code, never on the person.",
        "You must not engage in any external freelance or employment activity that "
        "conflicts with, competes with, or compromises your availability and commitment "
        "to your role at Codesino, without prior written approval from management.",
    ]:
        add_numbered(doc, item)

    if data.get("additionalNotice"):
        spacer(doc, 60, 40)
        heading2(doc, "5.6  Specific Instructions from Management")
        body_text(doc, data.get("additionalNotice", ""))

    divider(doc)

    # ─── ACCEPTANCE STATEMENT ─────────────────────────────────────────────
    heading1(doc, "6.  ACCEPTANCE OF EMPLOYMENT TERMS")
    spacer(doc, 40, 40)

    def acceptance_fn(cell):
        p = cell.add_paragraph()
        add_paragraph_spacing(p, before=0, after=60, line=276)
        add_run(p,
                "By accepting this employment and signing below, I confirm that:",
                bold=True, size_pt=11, color=BRAND_BLUE)
        for item in [
            "I have read and fully understood this Technical Employment Agreement in its entirety.",
            "I agree to abide by all security standards, code quality requirements, "
            "confidentiality obligations, and company policies outlined in this document.",
            "I understand that any breach of source code confidentiality, data security, "
            "or misconduct may result in immediate termination and legal prosecution.",
            "I accept the compensation, technical responsibilities, and conditions of "
            "employment as described in this agreement.",
            "I acknowledge that all code and technical work I produce during my employment "
            "is the exclusive intellectual property of Codesino Software Development Services.",
        ]:
            add_bullet(cell, item)
        pn = cell.add_paragraph()
        add_paragraph_spacing(pn, before=60, after=0)
        add_run(pn,
                "This agreement is effective from the date of hire specified in Section 2.",
                italic=True, size_pt=10, color=MUTED_BLUE)

    section_box(doc, acceptance_fn, HEX_LIGHT_BLUE, accent_color=HEX_MID_BLUE)
    divider(doc)

    _build_standard_clauses(doc)

    heading1(doc, "ADDITIONAL NOTES & SPECIAL CONDITIONS")
    notes = data.get("additionalNotes", "")
    if notes:
        body_text(doc, notes)
    else:
        italic_note(doc, "No additional notes for this agreement.")
    divider(doc)

    _build_signatures(doc, data)
    _build_footer(doc, "Technical Team Employment Agreement")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()